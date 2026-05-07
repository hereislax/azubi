"""
Views für die Kurs- und Praktikumsverwaltung.

Enthält Kurs-CRUD, Ablaufplan-Blöcke, Praktikumskalender, Einsatzverwaltung,
Zuweisungs-/Praktikumsplan-/Stationsschreiben, Kapazitätsplanung,
Kurskalender, Checklisten und Ausbildungsplan-Übersicht.
"""
# SPDX-License-Identifier: EUPL-1.2
# SPDX-FileCopyrightText: 2026 devNicoLax
from datetime import date, timedelta

from django.core.exceptions import PermissionDenied
from django.db import models
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.contrib import messages

from student.models import Student
from .forms import CourseForm, ScheduleBlockForm, InternshipAssignmentForm, SeminarLectureForm
from .models import (
    Course, ScheduleBlock, InternshipAssignment,
    BlockLetter, BlockLetterItem, BlockLetterTemplate,
    InternshipPlanLetter, InternshipPlanItem, InternshipPlanTemplate,
    StationLetter, StationLetterItem, StationLetterTemplate,
    SeminarLecture,
    ASSIGNMENT_STATUS_PENDING,
    BLOCK_LETTER_STATUS_PENDING, BLOCK_LETTER_STATUS_SENT,
    LECTURE_STATUS_PENDING, LECTURE_STATUS_CONFIRMED, LECTURE_STATUS_DECLINED,
    pick_active_letter_template,
)

from services.colors import BUNDESFARBEN_PALETTE as INTERNSHIP_COLORS  # noqa: E402


def _get_unit_capacity_info(schedule_block, exclude_assignment_pk=None, start_date=None, end_date=None):
    """
    Returns (full_pks, usage_map) for a given date range (or schedule block as fallback).

    full_pks  – set of unit PKs that cannot accept another assignment because
                they or an ancestor have reached max_capacity.
    usage_map – dict {unit_pk: (used, max_capacity)} for units that have a limit.

    When start_date/end_date are given, capacity is checked globally across ALL blocks
    for that date range — any two assignments that overlap in time compete for capacity,
    regardless of which block they belong to.
    When no dates are given, the check falls back to the given schedule_block only.
    """
    from organisation.models import OrganisationalUnit

    all_units = list(OrganisationalUnit.objects.all())

    # children_map: parent_pk → [child_pk, ...]
    children_map: dict[int, list[int]] = {}
    for u in all_units:
        children_map.setdefault(u.parent_id, [])
        children_map.setdefault(u.pk, [])
        if u.parent_id is not None:
            children_map[u.parent_id].append(u.pk)

    # Konkurrierende Einsätze für die Kapazitätsprüfung zählen
    if start_date and end_date:
        # Global: alle Blöcke, aber nur überlappendes Datumsfenster
        qs = InternshipAssignment.objects.filter(
            start_date__lte=end_date,
            end_date__gte=start_date,
        )
    else:
        # Fallback: gleicher Block, alle Daten
        qs = InternshipAssignment.objects.filter(schedule_block=schedule_block)
    if exclude_assignment_pk:
        qs = qs.exclude(pk=exclude_assignment_pk)
    direct_counts: dict[int, int] = {}
    for row in qs.values('unit_id'):
        direct_counts[row['unit_id']] = direct_counts.get(row['unit_id'], 0) + 1

    # Memoisierte rekursive Summe (Einheit + alle Nachfahren)
    memo: dict[int, int] = {}

    def total_count(pk: int) -> int:
        if pk in memo:
            return memo[pk]
        cnt = direct_counts.get(pk, 0)
        for child_pk in children_map.get(pk, []):
            cnt += total_count(child_pk)
        memo[pk] = cnt
        return cnt

    # Alle Einheiten sammeln, die ihr Limit erreicht haben – plus alle Nachfahren
    full_pks: set[int] = set()

    def mark_descendants(pk: int) -> None:
        full_pks.add(pk)
        for child_pk in children_map.get(pk, []):
            mark_descendants(child_pk)

    for u in all_units:
        if u.max_capacity is not None and total_count(u.pk) >= u.max_capacity:
            mark_descendants(u.pk)

    # Auslastungs-Map für die Anzeige (nur Einheiten mit Limit)
    usage_map = {
        u.pk: (total_count(u.pk), u.max_capacity)
        for u in all_units
        if u.max_capacity is not None
    }

    return full_pks, usage_map


def _build_months_data(period_start, period_end):
    """
    Berechnet die Monats-Kopfzeilen für den Gantt-Kursplan.

    Gibt eine Liste von Monatssegmenten zurück, die proportional zum übergebenen
    Datumsbereich skaliert sind. Jedes Segment enthält den Monatsnamen, eine
    Kurzform, die Anzahl der Tage im sichtbaren Ausschnitt sowie Offset und Breite
    als Prozentwerte (als Zeichenkette mit 6 Nachkommastellen) für die CSS-Positionierung.
    """
    period_days = (period_end - period_start).days + 1
    months_data = []
    current = period_start.replace(day=1)
    while current <= period_end:
        if current.month == 12:
            next_month = current.replace(year=current.year + 1, month=1, day=1)
        else:
            next_month = current.replace(month=current.month + 1, day=1)
        clip_start = max(current, period_start)
        clip_end = min(next_month - timedelta(days=1), period_end)
        m_days = (clip_end - clip_start).days + 1
        months_data.append({
            "name": current.strftime("%B"),
            "short": current.strftime("%b"),
            "days": m_days,
            "offset": f"{(clip_start - period_start).days / period_days * 100:.6f}",
            "width": f"{m_days / period_days * 100:.6f}",
        })
        current = next_month
    return months_data


@login_required()
def course_list(request):
    courses = Course.objects.select_related('job_profile').order_by('start_date')
    return render(request, 'course/course_list.html', {'courses': courses})


@login_required()
def course_detail(request, pk):
    from datetime import date
    from services.paperless import PaperlessService
    from services.roles import is_training_director, is_training_office
    from student.models import ChecklistTemplate, StudentChecklist
    from course.models import CourseChecklist
    course = get_object_or_404(Course.objects.select_related('job_profile__career', 'job_profile__specialization'), pk=pk)
    students = Student.objects.filter(course=course).select_related('gender', 'status').order_by('last_name', 'first_name')
    schedule_blocks = course.schedule_blocks.prefetch_related('letters', 'plan_letters', 'station_letters').all()
    can_view_akte = is_training_director(request.user) or is_training_office(request.user)
    can_manage_checklists = is_training_director(request.user) or is_training_office(request.user)
    paperless_docs = PaperlessService.get_documents_for_course(course.title) if can_view_akte else None

    # Checklisten-Übersicht pro NK als Template-freundliche Liste
    checklist_templates = ChecklistTemplate.objects.filter(is_active=True) if can_manage_checklists else []
    if can_manage_checklists and course.job_profile:
        # Vorlagen für das Berufsbild des Kurses oder universelle Vorlagen (keine Zuordnung)
        checklist_templates = ChecklistTemplate.objects.filter(is_active=True).filter(
            models.Q(job_profiles=course.job_profile) | models.Q(job_profiles__isnull=True)
        ).distinct()
    students_with_checklists = []
    if can_manage_checklists:
        student_pks = [s.pk for s in students]
        all_cls = (
            StudentChecklist.objects
            .filter(student_id__in=student_pks)
            .prefetch_related('items')
            .select_related('template')
            .order_by('template__name', 'name')
        )
        cl_map = {}
        for cl in all_cls:
            cl_map.setdefault(cl.student_id, []).append(cl)
        for s in students:
            students_with_checklists.append({
                'student': s,
                'checklists': cl_map.get(s.pk, []),
            })

    # Kurs-Checklisten
    course_checklists = (
        CourseChecklist.objects
        .filter(course=course)
        .prefetch_related('items')
        .select_related('template')
        .order_by('name')
    ) if can_manage_checklists else []

    return render(request, 'course/course_detail.html', {
        'course': course,
        'students': students,
        'schedule_blocks': schedule_blocks,
        'today': date.today(),
        'can_view_akte': can_view_akte,
        'paperless_docs': paperless_docs,
        'can_manage_checklists': can_manage_checklists,
        'checklist_templates': checklist_templates,
        'students_with_checklists': students_with_checklists,
        'course_checklists': course_checklists,
    })


@login_required
def course_competence_matrix(request, pk):
    """Aggregierte Kompetenzmatrix aller Nachwuchskräfte eines Kurses (Heatmap)."""
    from django.core.exceptions import PermissionDenied
    from services.roles import (
        is_training_director, is_training_office,
        is_training_coordinator, get_chief_instructor,
    )
    from services.competence_matrix import get_competence_matrix
    from student.models import Student
    from organisation.models import Competence
    from course.models import CompetenceTarget
    from instructor.views import _get_coordination_area

    course = get_object_or_404(Course.objects.select_related('job_profile'), pk=pk)
    user = request.user
    can_view_full = (
        user.is_staff
        or is_training_director(user)
        or is_training_office(user)
    )

    coord_unit_pks = None
    if not can_view_full:
        if not is_training_coordinator(user):
            raise PermissionDenied
        chief = get_chief_instructor(user)
        if not (chief and chief.coordination):
            raise PermissionDenied
        coord_unit_pks, _, _ = _get_coordination_area(chief.coordination)

    students = list(
        Student.objects
        .filter(course=course, anonymized_at__isnull=True)
        .select_related('course__job_profile')
        .order_by('last_name', 'first_name')
    )

    # Koordination: nur Azubis, die bei ihr eingesetzt sind/waren
    if coord_unit_pks is not None:
        from datetime import timedelta
        cutoff = date.today() - timedelta(days=14)
        student_ids_with_access = set(
            InternshipAssignment.objects
            .filter(student__course=course, unit_id__in=coord_unit_pks, end_date__gte=cutoff)
            .values_list('student_id', flat=True).distinct()
        )
        students = [s for s in students if s.pk in student_ids_with_access]

    # Kompetenz-Spalten: alle, die Endziele oder Mappings für dieses Berufsbild haben
    job_profile = course.job_profile
    if job_profile is None:
        return render(request, 'course/competence_matrix.html', {
            'course': course, 'students_rows': [], 'competences': [], 'no_data': True,
        })

    target_competence_ids = set(
        CompetenceTarget.objects.filter(job_profile=job_profile).values_list('competence_id', flat=True)
    )
    from assessment.models import CriterionCompetenceWeight
    mapping_competence_ids = set(
        CriterionCompetenceWeight.objects
        .filter(criterion__job_profile=job_profile)
        .values_list('competence_id', flat=True)
    )
    competence_ids = target_competence_ids | mapping_competence_ids
    competences = list(
        Competence.objects.filter(pk__in=competence_ids).order_by('name')
    )

    students_rows = []
    for s in students:
        m = get_competence_matrix(s)
        comp_lookup = {c['competence'].pk: c for c in m['competences']}
        cells = []
        for comp in competences:
            row = comp_lookup.get(comp.pk)
            if row is None:
                cells.append({'value': None, 'self_value': None, 'target': None, 'coverage': None})
            else:
                cells.append({
                    'value':       row['external_value'],
                    'self_value':  row['self_value'],
                    'target':      row['target_now'],
                    'coverage':    row['coverage_pct'],
                })
        students_rows.append({
            'student':                 s,
            'cells':                   cells,
            'apprentice_progress_pct': m['apprentice_progress_pct'],
        })

    return render(request, 'course/competence_matrix.html', {
        'course':        course,
        'job_profile':   job_profile,
        'competences':   competences,
        'students_rows': students_rows,
        'no_data':       not (competences and students_rows),
    })


@login_required
@require_POST
def course_document_upload(request, pk):
    from django.core.exceptions import PermissionDenied
    from services.roles import is_training_director, is_training_office
    from services.paperless import PaperlessService
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied
    course = get_object_or_404(Course, pk=pk)
    uploaded_file = request.FILES.get('file')
    title = request.POST.get('title', '').strip()
    if not uploaded_file or not title:
        messages.error(request, 'Bitte Titel und Datei angeben.')
        return redirect(f'/kurs/{pk}/?tab=akte')
    from services.validators import validate_document
    from django.core.exceptions import ValidationError as DjangoValidationError
    try:
        validate_document(uploaded_file)
    except DjangoValidationError as e:
        messages.error(request, str(e.message))
        return redirect(f'/kurs/{pk}/?tab=akte')
    doc_id = PaperlessService.upload_and_wait_for_course(
        file_bytes=uploaded_file.read(),
        title=title,
        course_title=course.title,
        filename=uploaded_file.name,
        mime_type=uploaded_file.content_type or 'application/octet-stream',
    )
    if doc_id:
        messages.success(request, f'„{title}" wurde erfolgreich in die Kursakte hochgeladen.')
    else:
        messages.error(request, 'Upload zu Paperless fehlgeschlagen. Bitte erneut versuchen.')
    return redirect(f'/kurs/{pk}/?tab=akte')


@login_required
def course_create(request):
    form = CourseForm(request.POST or None)
    if form.is_valid():
        course = form.save()
        messages.success(request, f'Kurs „{course.title}" wurde erfolgreich angelegt.')
        return redirect('course:course_detail', pk=course.pk)
    return render(request, 'course/course_form.html', {'form': form, 'action': 'Anlegen'})


@login_required
def course_edit(request, pk):
    course = get_object_or_404(Course, pk=pk)
    form = CourseForm(request.POST or None, instance=course)
    if form.is_valid():
        form.save()
        messages.success(request, f'Kurs „{course.title}" wurde erfolgreich gespeichert.')
        return redirect('course:course_detail', pk=course.pk)
    return render(request, 'course/course_form.html', {
        'form': form,
        'action': 'Bearbeiten',
        'course': course,
    })


@login_required
def schedule_block_create(request, course_pk):
    course = get_object_or_404(Course, pk=course_pk)
    form = ScheduleBlockForm(request.POST or None, course=course)
    if form.is_valid():
        block = form.save(commit=False)
        block.course = course
        block.save()
        messages.success(request, f'Block „{block.name}" wurde angelegt.')
        return redirect('course:course_detail', pk=course.pk)
    return render(request, 'course/schedule_block_form.html', {
        'form': form,
        'course': course,
        'action': 'Anlegen',
    })


@login_required
def schedule_block_edit(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    form = ScheduleBlockForm(request.POST or None, instance=block, course=course)
    if form.is_valid():
        form.save()
        messages.success(request, f'Block „{block.name}" wurde gespeichert.')
        return redirect('course:course_detail', pk=course.pk)
    return render(request, 'course/schedule_block_form.html', {
        'form': form,
        'course': course,
        'schedule_block': block,
        'action': 'Bearbeiten',
    })


@login_required
def schedule_block_delete(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    if request.method == 'POST':
        name = block.name
        block.delete()
        messages.success(request, f'Block „{name}" wurde gelöscht.')
        return redirect('course:course_detail', pk=course.pk)
    return render(request, 'course/schedule_block_confirm_delete.html', {
        'schedule_block': block,
        'course': course,
    })


@login_required
def course_delete(request, pk):
    course = get_object_or_404(Course, pk=pk)
    if request.method == 'POST':
        title = course.title
        try:
            course.delete()
            messages.success(request, f'Kurs „{title}" wurde gelöscht.')
        except Exception:
            messages.error(request, f'Kurs „{title}" kann nicht gelöscht werden, da noch Nachwuchskräfte zugeordnet sind.')
            return redirect('course:course_detail', pk=pk)
        return redirect('course:course_list')
    return render(request, 'course/course_confirm_delete.html', {'course': course})


# ── Internship views ────────────────────────────────────────────────────────
@login_required
def internship_calendar(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')

    block_start = block.start_date
    block_end = block.end_date
    block_days = (block_end - block_start).days + 1

    months_data = _build_months_data(block_start, block_end)

    today = date.today()
    today_offset = None
    if block_start <= today <= block_end:
        today_offset = f"{(today - block_start).days / block_days * 100:.4f}"

    # Konsistente Farbe pro Einheit (anhand pk)
    from organisation.models import OrganisationalUnit
    all_units = {u.pk: u for u in OrganisationalUnit.objects.all()}
    unit_color = lambda uid: INTERNSHIP_COLORS[uid % len(INTERNSHIP_COLORS)]

    students = Student.objects.filter(course=course).select_related('gender').order_by('last_name', 'first_name')

    bar_height = 55
    bar_gap = 6
    row_padding = 6

    rows_data = []
    used_unit_pks = set()

    for student in students:
        assignments_qs = InternshipAssignment.objects.filter(
            schedule_block=block, student=student,
        ).select_related('unit').order_by('start_date')

        assignments_data = []
        for a in assignments_qs:
            clip_start = max(a.start_date, block_start)
            clip_end = min(a.end_date, block_end)
            if clip_start > clip_end:
                continue
            raw_offset = (clip_start - block_start).days / block_days * 100
            raw_width = max(((clip_end - clip_start).days + 1) / block_days * 100, 0.3)
            used_unit_pks.add(a.unit_id)
            assignments_data.append({
                "assignment": a,
                "offset": f"{raw_offset:.4f}",
                "width": f"{raw_width:.4f}",
                "raw_offset": raw_offset,
                "raw_end": raw_offset + raw_width,
                "label": a.unit.name,
                "start_fmt": a.start_date.strftime("%d.%m.%Y"),
                "end_fmt": a.end_date.strftime("%d.%m.%Y"),
            })

        # Nachbarschaftsbewusste Farbzuweisung (sortiert nach start_date):
        # Bevorzugt die Hash-basierte Farbe der Einheit, vermeidet aber Farben,
        # die ein überlappender oder direkt benachbarter Balken bereits trägt.
        for i, a_data in enumerate(assignments_data):
            preferred = a_data["assignment"].unit_id % len(INTERNSHIP_COLORS)
            used_colors = {
                other["color_idx"]
                for other in assignments_data[:i]
                if a_data["raw_offset"] <= other["raw_end"]
                and a_data["raw_end"] >= other["raw_offset"]
            }
            color_idx = preferred
            while color_idx in used_colors:
                color_idx = (color_idx + 1) % len(INTERNSHIP_COLORS)
            a_data["color_idx"] = color_idx
            a_data["color"] = INTERNSHIP_COLORS[color_idx]

        # Lane packing
        lane_ends = []
        for a_data in assignments_data:
            placed = False
            for i, end in enumerate(lane_ends):
                if a_data["raw_offset"] >= end:
                    lane_ends[i] = a_data["raw_end"]
                    a_data["lane"] = i
                    placed = True
                    break
            if not placed:
                a_data["lane"] = len(lane_ends)
                lane_ends.append(a_data["raw_end"])

        num_lanes = max(len(lane_ends), 1)
        row_height = bar_gap + num_lanes * (bar_height + bar_gap) + row_padding
        for a_data in assignments_data:
            a_data["top_px"] = bar_gap + a_data["lane"] * (bar_height + bar_gap)

        rows_data.append({
            "student": student,
            "assignments": assignments_data,
            "row_height": row_height,
        })

    units_legend = [
        {"unit": all_units[pk], "color": unit_color(pk)}
        for pk in sorted(used_unit_pks)
        if pk in all_units
    ]

    return render(request, 'course/internship_calendar.html', {
        'course': course,
        'schedule_block': block,
        'block_start_iso': block_start.isoformat(),
        'block_days': block_days,
        'months_data': months_data,
        'today_offset': today_offset,
        'rows_data': rows_data,
        'units_legend': units_legend,
    })


@login_required
def internship_assignment_create(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    initial = {k: request.GET[k] for k in ('start_date', 'end_date', 'student', 'unit') if k in request.GET}
    from datetime import datetime as _dt
    def _parse_date(s):
        try:
            return _dt.strptime(s, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return None
    start_date = _parse_date(request.GET.get('start_date'))
    end_date   = _parse_date(request.GET.get('end_date'))
    full_pks, usage_map = _get_unit_capacity_info(block, start_date=start_date, end_date=end_date)
    form = InternshipAssignmentForm(request.POST or None, block=block, course=course,
                                    initial=initial, full_unit_pks=full_pks)
    if form.is_valid():
        assignment = form.save(commit=False)
        assignment.schedule_block = block
        assignment.created_by = request.user
        assignment.save()
        from services.notifications import notify_chiefs_of_assignment, notify_instructor_of_assignment
        notify_chiefs_of_assignment(request, assignment, is_new=True)
        notify_instructor_of_assignment(request, assignment)
        messages.success(request, f'Praktikumseinsatz für {assignment.student} wurde angelegt.')
        return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
    import json
    return render(request, 'course/internship_assignment_form.html', {
        'form': form,
        'course': course,
        'schedule_block': block,
        'action': 'Anlegen',
        'full_unit_pks_json': json.dumps(list(full_pks)),
        'usage_map_json': json.dumps({str(k): v for k, v in usage_map.items()}),
        'job_profile_pk': course.job_profile_id,
    })


@login_required
def internship_assignment_edit(request, course_pk, block_public_id, assignment_pk):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    assignment = get_object_or_404(InternshipAssignment, pk=assignment_pk, schedule_block=block)
    previous_instructor_id = assignment.instructor_id
    full_pks, usage_map = _get_unit_capacity_info(block, exclude_assignment_pk=assignment.pk,
                                                   start_date=assignment.start_date, end_date=assignment.end_date)
    form = InternshipAssignmentForm(request.POST or None, instance=assignment, block=block,
                                    course=course, full_unit_pks=full_pks)
    if form.is_valid():
        updated = form.save(commit=False)
        updated.status = ASSIGNMENT_STATUS_PENDING
        updated.rejection_reason = ''
        updated.save()
        assignment.refresh_from_db()
        if assignment.instructor_id and assignment.instructor_id != previous_instructor_id:
            assignment.bump_notification_sequence()
            from services.notifications import notify_instructor_of_assignment
            notify_instructor_of_assignment(request, assignment)
        messages.success(request, f'Praktikumseinsatz für {assignment.student} wurde gespeichert.')
        return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
    import json
    return render(request, 'course/internship_assignment_form.html', {
        'form': form,
        'course': course,
        'schedule_block': block,
        'assignment': assignment,
        'action': 'Bearbeiten',
        'full_unit_pks_json': json.dumps(list(full_pks)),
        'usage_map_json': json.dumps({str(k): v for k, v in usage_map.items()}),
        'job_profile_pk': course.job_profile_id,
    })


@login_required
def internship_assignment_delete(request, course_pk, block_public_id, assignment_pk):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    assignment = get_object_or_404(InternshipAssignment, pk=assignment_pk, schedule_block=block)
    if request.method == 'POST':
        label = str(assignment)
        assignment.delete()
        messages.success(request, f'Einsatz „{label}" wurde gelöscht.')
        return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
    return render(request, 'course/internship_assignment_confirm_delete.html', {
        'assignment': assignment,
        'course': course,
        'schedule_block': block,
    })


# ── Zuweisungsschreiben ──────────────────────────────────────────────────────

@login_required
def block_letter_create(request, course_pk, block_public_id):
    """Formular: Freitext eingeben und Zuweisungsschreiben-Stapel starten."""
    from django.core.exceptions import PermissionDenied
    from services.roles import is_training_director, is_training_office

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)

    if not (request.user.is_staff or is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied

    students = Student.objects.filter(course=course).order_by('last_name', 'first_name')
    template_obj = pick_active_letter_template(BlockLetterTemplate, course.job_profile)

    if request.method == 'POST':
        free_text = request.POST.get('free_text', '')
        if not template_obj:
            messages.error(request, 'Es ist keine aktive Zuweisungsschreiben-Vorlage hinterlegt. Bitte zuerst eine Vorlage im Admin hochladen.')
            return redirect('course:course_detail', pk=course.pk)
        if not students.exists():
            messages.error(request, 'Dem Kurs sind keine Nachwuchskräfte zugeordnet.')
            return redirect('course:course_detail', pk=course.pk)

        from django.utils import timezone as tz
        is_leitung = request.user.is_staff or is_training_director(request.user)
        letter = BlockLetter.objects.create(
            schedule_block=block,
            free_text=free_text,
            status=BLOCK_LETTER_STATUS_SENT if is_leitung else BLOCK_LETTER_STATUS_PENDING,
            generated_by=request.user,
            approved_by=request.user if is_leitung else None,
            approved_at=tz.now() if is_leitung else None,
        )
        if not is_leitung:
            from django.urls import reverse
            from services.models import notify_staff
            notify_staff(
                message=f'Zuweisungsschreiben zur Freigabe: {block.name} ({course.pk})',
                link=reverse('course:block_letter_approve', kwargs={'course_pk': course.pk, 'block_public_id': block.public_id, 'letter_pk': letter.pk}),
                icon='bi-envelope-check',
                category='Schreiben',
            )
        return redirect('course:block_letter_generate', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    return render(request, 'course/block_letter_form.html', {
        'course': course,
        'schedule_block': block,
        'students': students,
        'has_template': template_obj is not None,
    })


@login_required
def block_letter_generate(request, course_pk, block_public_id, letter_pk):
    """
    Ladeseite: GET zeigt Spinner, JS feuert sofort einen AJAX-POST.
    POST generiert alle Schreiben und gibt JSON zurück.
    """
    import logging
    from datetime import date as date_
    from io import BytesIO
    from django.http import JsonResponse
    from django.core.exceptions import PermissionDenied
    from services.roles import is_training_director, is_training_office
    from services.paperless import PaperlessService
    from services.email import send_mail

    logger = logging.getLogger(__name__)

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    letter = get_object_or_404(BlockLetter, pk=letter_pk, schedule_block=block)

    if not (request.user.is_staff or is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied

    if request.method == 'GET':
        return render(request, 'course/block_letter_loading.html', {
            'course': course,
            'schedule_block': block,
            'letter': letter,
        })

    # POST – Verarbeitung durchführen
    template_obj = pick_active_letter_template(BlockLetterTemplate, course.job_profile)
    if not template_obj:
        return JsonResponse({'error': 'Keine aktive Vorlage gefunden.'})

    students = Student.objects.filter(course=course).select_related('gender', 'address').order_by('last_name', 'first_name')
    creator = letter.generated_by
    profile = getattr(creator, 'profile', None)

    send_immediately = letter.status == BLOCK_LETTER_STATUS_SENT

    if not send_immediately:
        # Dokumente werden erst bei der Freigabe mit dem Namen der genehmigenden Person generiert
        detail_url = f'/kurs/{course.pk}/ablaufplan/{block.public_id}/zuweisungsschreiben/{letter.pk}/'
        return JsonResponse({'redirect_url': detail_url})

    errors = []

    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    for student in students:
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **creator_context(creator),
                **meta_context(),
                'freitext': letter.free_text,
                'zeichnung': letter.approved_by.last_name if letter.approved_by else 'ohne Zeichnung',
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Zuweisungsschreiben: Vorlage für %s konnte nicht gerendert werden: %s', student, exc)
            errors.append(str(student))
            continue

        title = f'Zuweisungsschreiben {block.name} – {student.first_name} {student.last_name}'
        filename = f'zuweisungsschreiben_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
        doc_id = upload_to_paperless(
            file_bytes=file_bytes,
            title=title,
            student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Zuweisungsschreiben: Paperless-Upload für %s fehlgeschlagen', student)
            errors.append(str(student))
            continue

        item = BlockLetterItem.objects.create(letter=letter, student=student, paperless_id=doc_id)

        if send_immediately and student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            attachments = [(f'{filename[:-5]}.pdf', pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Zuweisungsschreiben: {block.name}',
                    body_text=f'Guten Tag {student.first_name} {student.last_name},\n\nim Anhang erhalten Sie Ihr Zuweisungsschreiben für den Block „{block.name}".\n\nMit freundlichen Grüßen\nIhr Azubi-Portal',
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Zuweisungsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    detail_url = (
        f'/kurs/{course.pk}/ablaufplan/{block.public_id}/zuweisungsschreiben/{letter.pk}/'
    )
    if errors:
        return JsonResponse({'error': f'Fehler bei folgenden Nachwuchskräften: {", ".join(errors)}. Teilweise wurden Schreiben generiert.', 'redirect_url': detail_url})
    return JsonResponse({'redirect_url': detail_url})


@login_required
def block_letter_detail(request, course_pk, block_public_id, letter_pk):
    """Übersicht über einen Zuweisungsschreiben-Stapel."""
    from services.roles import is_training_director, is_training_office
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    letter = get_object_or_404(BlockLetter, pk=letter_pk, schedule_block=block)
    items = letter.items.select_related('student__gender').order_by('student__last_name', 'student__first_name')
    return render(request, 'course/block_letter_detail.html', {
        'course': course,
        'schedule_block': block,
        'letter': letter,
        'items': items,
        'can_regenerate': _is_leitung_or_referat(request.user),
    })


@login_required
def block_letter_approve(request, course_pk, block_public_id, letter_pk):
    """Ausbildungsleitung gibt ausstehende Zuweisungsschreiben frei und versendet sie."""
    import logging
    from django.http import JsonResponse
    from django.core.exceptions import PermissionDenied
    from django.utils import timezone
    from services.roles import is_training_director
    from services.paperless import PaperlessService
    from services.email import send_mail

    logger = logging.getLogger(__name__)

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    letter = get_object_or_404(BlockLetter, pk=letter_pk, schedule_block=block)

    if not (request.user.is_staff or is_training_director(request.user)):
        raise PermissionDenied

    if request.method == 'GET':
        students = Student.objects.filter(course=course).select_related('gender').order_by('last_name', 'first_name')
        return render(request, 'course/block_letter_approve.html', {
            'course': course,
            'schedule_block': block,
            'letter': letter,
            'students': students,
        })

    # POST – Dokumente mit Unterzeichner generieren, in Paperless ablegen, versenden
    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    letter.approved_by = request.user
    letter.approved_at = timezone.now()
    letter.save(update_fields=['approved_by', 'approved_at'])

    template_obj = pick_active_letter_template(BlockLetterTemplate, course.job_profile)
    if not template_obj:
        messages.error(request, 'Keine aktive Vorlage gefunden. Freigabe nicht möglich.')
        return redirect('course:block_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    students = Student.objects.filter(course=course).select_related('gender', 'address').order_by('last_name', 'first_name')
    creator = letter.generated_by
    errors = []

    for student in students:
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **creator_context(creator),
                **meta_context(),
                'freitext': letter.free_text,
                'zeichnung': letter.approved_by.last_name,
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Freigabe Zuweisungsschreiben: Rendering für %s fehlgeschlagen: %s', student, exc)
            errors.append(str(student))
            continue

        title = f'Zuweisungsschreiben {block.name} – {student.first_name} {student.last_name}'
        filename = f'zuweisungsschreiben_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
        doc_id = upload_to_paperless(
            file_bytes=file_bytes,
            title=title,
            student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Freigabe Zuweisungsschreiben: Upload für %s fehlgeschlagen', student)
            errors.append(str(student))
            continue

        item, _ = BlockLetterItem.objects.get_or_create(letter=letter, student=student)
        if item.paperless_id and item.paperless_id != doc_id:
            PaperlessService.delete_document(item.paperless_id)
        item.paperless_id = doc_id
        item.email_sent = False
        item.save(update_fields=['paperless_id', 'email_sent'])

        if student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            filename_pdf = f'zuweisungsschreiben_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.pdf'
            attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Zuweisungsschreiben: {block.name}',
                    body_text=f'Guten Tag {student.first_name} {student.last_name},\n\nim Anhang erhalten Sie Ihr Zuweisungsschreiben für den Block „{block.name}".\n\nMit freundlichen Grüßen\nIhr Azubi-Portal',
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Freigabe Zuweisungsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)
                errors.append(str(student))

    letter.status = BLOCK_LETTER_STATUS_SENT
    letter.save(update_fields=['status'])

    if errors:
        messages.warning(request, f'Freigabe erteilt, aber Fehler bei folgenden Nachwuchskräften: {", ".join(errors)}')
    else:
        messages.success(request, f'Zuweisungsschreiben für „{block.name}" freigegeben und an alle Nachwuchskräfte versendet.')
    return redirect('course:block_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


# ── Praktikumspläne ──────────────────────────────────────────────────────────

def _is_leitung_or_referat(user):
    from services.roles import is_training_director, is_training_office
    return user.is_staff or is_training_director(user) or is_training_office(user)


@login_required
def internship_plan_create(request, course_pk, block_public_id):
    from django.core.exceptions import PermissionDenied
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    students = Student.objects.filter(course=course).order_by('last_name', 'first_name')
    template_obj = pick_active_letter_template(InternshipPlanTemplate, course.job_profile)

    if request.method == 'POST':
        free_text = request.POST.get('free_text', '')
        if not template_obj:
            messages.error(request, 'Keine aktive Praktikumsplan-Vorlage hinterlegt.')
            return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
        if not students.exists():
            messages.error(request, 'Dem Kurs sind keine Nachwuchskräfte zugeordnet.')
            return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
        from django.utils import timezone as tz
        from services.roles import is_training_director
        is_leitung = request.user.is_staff or is_training_director(request.user)
        letter = InternshipPlanLetter.objects.create(
            schedule_block=block,
            free_text=free_text,
            status=BLOCK_LETTER_STATUS_SENT if is_leitung else BLOCK_LETTER_STATUS_PENDING,
            generated_by=request.user,
            approved_by=request.user if is_leitung else None,
            approved_at=tz.now() if is_leitung else None,
        )
        if not is_leitung:
            from django.urls import reverse
            from services.models import notify_staff
            notify_staff(
                message=f'Praktikumsplan zur Freigabe: {block.name} ({course.pk})',
                link=reverse('course:internship_plan_approve', kwargs={'course_pk': course.pk, 'block_public_id': block.public_id, 'letter_pk': letter.pk}),
                icon='bi-map',
                category='Schreiben',
            )
        return redirect('course:internship_plan_generate', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    return render(request, 'course/internship_plan_form.html', {
        'course': course,
        'schedule_block': block,
        'students': students,
        'has_template': template_obj is not None,
    })


@login_required
def internship_plan_generate(request, course_pk, block_public_id, letter_pk):
    import logging
    from django.http import JsonResponse
    from django.core.exceptions import PermissionDenied
    from services.paperless import PaperlessService
    from services.email import send_mail
    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(InternshipPlanLetter, pk=letter_pk, schedule_block=block)

    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    if request.method == 'GET':
        return render(request, 'course/internship_plan_loading.html', {
            'course': course, 'schedule_block': block, 'letter': letter,
        })

    template_obj = pick_active_letter_template(InternshipPlanTemplate, course.job_profile)
    if not template_obj:
        return JsonResponse({'error': 'Keine aktive Vorlage gefunden.'})

    students = Student.objects.filter(course=course).select_related('gender', 'address').order_by('last_name', 'first_name')
    creator = letter.generated_by
    send_immediately = letter.status == BLOCK_LETTER_STATUS_SENT

    if not send_immediately:
        # Dokumente werden erst bei der Freigabe mit dem Namen der genehmigenden Person generiert
        detail_url = f'/kurs/{course.pk}/ablaufplan/{block.public_id}/praktikumsplan/{letter.pk}/'
        return JsonResponse({'redirect_url': detail_url})

    errors = []

    for student in students:
        assignments = (
            InternshipAssignment.objects
            .filter(schedule_block=block, student=student)
            .select_related('unit', 'location', 'instructor')
            .order_by('start_date')
        )
        einsaetze = [
            {
                'einheit': a.unit.name,
                'beginn': a.start_date.strftime('%d.%m.%Y'),
                'ende': a.end_date.strftime('%d.%m.%Y'),
                'standort': str(a.location) if a.location else '',
                'praxistutor': (
                    f'{a.instructor.first_name} {a.instructor.last_name}'
                    if a.instructor else ''
                ),
            }
            for a in assignments
        ]
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **creator_context(creator),
                **meta_context(),
                'anrede': (
                    f'{student.gender.description} {student.last_name},'
                    if student.gender else f'{student.first_name} {student.last_name},'
                ),
                'freitext': letter.free_text,
                'einsaetze': einsaetze,
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Praktikumsplan: Rendering für %s fehlgeschlagen: %s', student, exc)
            errors.append(str(student))
            continue

        title = f'Praktikumsplan {block.name} – {student.first_name} {student.last_name}'
        filename = f'praktikumsplan_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
        doc_id = upload_to_paperless(
            file_bytes=file_bytes, title=title, student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Praktikumsplan: Upload für %s fehlgeschlagen', student)
            errors.append(str(student))
            continue

        item = InternshipPlanItem.objects.create(letter=letter, student=student, paperless_id=doc_id)

        if send_immediately and student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            attachments = [(f'{filename[:-5]}.pdf', pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Praktikumsplan: {block.name}',
                    body_text=(
                        f'Guten Tag {student.first_name} {student.last_name},\n\n'
                        f'im Anhang erhalten Sie Ihren Praktikumsplan für den Block „{block.name}".\n\n'
                        f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                    ),
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Praktikumsplan: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    detail_url = f'/kurs/{course.pk}/ablaufplan/{block.public_id}/praktikumsplan/{letter.pk}/'
    if errors:
        return JsonResponse({'error': f'Fehler bei: {", ".join(errors)}', 'redirect_url': detail_url})
    return JsonResponse({'redirect_url': detail_url})


@login_required
def internship_plan_detail(request, course_pk, block_public_id, letter_pk):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(InternshipPlanLetter, pk=letter_pk, schedule_block=block)
    items = letter.items.select_related('student__gender').order_by('student__last_name', 'student__first_name')
    return render(request, 'course/internship_plan_detail.html', {
        'course': course, 'schedule_block': block, 'letter': letter, 'items': items,
        'can_regenerate': _is_leitung_or_referat(request.user),
    })


@login_required
def internship_plan_approve(request, course_pk, block_public_id, letter_pk):
    import logging
    from django.core.exceptions import PermissionDenied
    from django.utils import timezone
    from services.roles import is_training_director
    from services.paperless import PaperlessService
    from services.email import send_mail

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(InternshipPlanLetter, pk=letter_pk, schedule_block=block)

    if not (request.user.is_staff or is_training_director(request.user)):
        raise PermissionDenied

    if request.method == 'GET':
        students = Student.objects.filter(course=course).select_related('gender').order_by('last_name', 'first_name')
        return render(request, 'course/internship_plan_approve.html', {
            'course': course, 'schedule_block': block, 'letter': letter, 'students': students,
        })

    # POST – Dokumente mit Unterzeichner generieren, in Paperless ablegen, versenden
    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    letter.approved_by = request.user
    letter.approved_at = timezone.now()
    letter.save(update_fields=['approved_by', 'approved_at'])

    template_obj = pick_active_letter_template(InternshipPlanTemplate, course.job_profile)
    if not template_obj:
        messages.error(request, 'Keine aktive Vorlage gefunden. Freigabe nicht möglich.')
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    students = Student.objects.filter(course=course).select_related('gender', 'address').order_by('last_name', 'first_name')
    creator = letter.generated_by
    errors = []

    for student in students:
        assignments = (
            InternshipAssignment.objects
            .filter(schedule_block=block, student=student)
            .select_related('unit', 'location', 'instructor')
            .order_by('start_date')
        )
        einsaetze = [
            {
                'einheit': a.unit.name,
                'beginn': a.start_date.strftime('%d.%m.%Y'),
                'ende': a.end_date.strftime('%d.%m.%Y'),
                'standort': str(a.location) if a.location else '',
                'praxistutor': (
                    f'{a.instructor.first_name} {a.instructor.last_name}'
                    if a.instructor else ''
                ),
            }
            for a in assignments
        ]
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **creator_context(creator),
                **meta_context(),
                'anrede': (
                    f'{student.gender.description} {student.last_name},'
                    if student.gender else f'{student.first_name} {student.last_name},'
                ),
                'freitext': letter.free_text,
                'einsaetze': einsaetze,
                'zeichnung': letter.approved_by.last_name,
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Freigabe Praktikumsplan: Rendering für %s fehlgeschlagen: %s', student, exc)
            errors.append(str(student))
            continue

        title = f'Praktikumsplan {block.name} – {student.first_name} {student.last_name}'
        filename = f'praktikumsplan_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
        doc_id = upload_to_paperless(
            file_bytes=file_bytes, title=title, student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Freigabe Praktikumsplan: Upload für %s fehlgeschlagen', student)
            errors.append(str(student))
            continue

        item, _ = InternshipPlanItem.objects.get_or_create(letter=letter, student=student)
        if item.paperless_id and item.paperless_id != doc_id:
            PaperlessService.delete_document(item.paperless_id)
        item.paperless_id = doc_id
        item.email_sent = False
        item.save(update_fields=['paperless_id', 'email_sent'])

        if student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            filename_pdf = f'praktikumsplan_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.pdf'
            attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Praktikumsplan: {block.name}',
                    body_text=(
                        f'Guten Tag {student.first_name} {student.last_name},\n\n'
                        f'im Anhang erhalten Sie Ihren Praktikumsplan für den Block „{block.name}".\n\n'
                        f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                    ),
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Freigabe Praktikumsplan: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)
                errors.append(str(student))

    letter.status = BLOCK_LETTER_STATUS_SENT
    letter.save(update_fields=['status'])

    if errors:
        messages.warning(request, f'Freigabe erteilt, aber Fehler bei: {", ".join(errors)}')
    else:
        messages.success(request, f'Praktikumspläne für „{block.name}" freigegeben und versendet.')
    return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


# ── Stationszuweisungsschreiben ──────────────────────────────────────────────

@login_required
def station_letter_create(request, course_pk, block_public_id):
    from django.core.exceptions import PermissionDenied
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    assignments = (
        InternshipAssignment.objects
        .filter(schedule_block=block)
        .select_related('student__gender', 'unit')
        .order_by('student__last_name', 'student__first_name', 'start_date')
    )
    template_obj = pick_active_letter_template(StationLetterTemplate, course.job_profile)

    if request.method == 'POST':
        free_text = request.POST.get('free_text', '')
        if not template_obj:
            messages.error(request, 'Keine aktive Stationsschreiben-Vorlage hinterlegt.')
            return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
        if not assignments.exists():
            messages.error(request, 'Es sind noch keine Einsätze für diesen Block angelegt.')
            return redirect('course:internship_calendar', course_pk=course.pk, block_public_id=block.public_id)
        from django.utils import timezone as tz
        from services.roles import is_training_director
        is_leitung = request.user.is_staff or is_training_director(request.user)
        letter = StationLetter.objects.create(
            schedule_block=block,
            free_text=free_text,
            status=BLOCK_LETTER_STATUS_SENT if is_leitung else BLOCK_LETTER_STATUS_PENDING,
            generated_by=request.user,
            approved_by=request.user if is_leitung else None,
            approved_at=tz.now() if is_leitung else None,
        )
        if not is_leitung:
            from django.urls import reverse
            from services.models import notify_staff
            notify_staff(
                message=f'Stationsschreiben zur Freigabe: {block.name} ({course.pk})',
                link=reverse('course:station_letter_approve', kwargs={'course_pk': course.pk, 'block_public_id': block.public_id, 'letter_pk': letter.pk}),
                icon='bi-building-check',
                category='Schreiben',
            )
        return redirect('course:station_letter_generate', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    return render(request, 'course/station_letter_form.html', {
        'course': course,
        'schedule_block': block,
        'assignments': assignments,
        'has_template': template_obj is not None,
    })


@login_required
def station_letter_generate(request, course_pk, block_public_id, letter_pk):
    import logging
    from datetime import date as date_
    from io import BytesIO
    from django.http import JsonResponse
    from django.core.exceptions import PermissionDenied
    from services.paperless import PaperlessService
    from services.email import send_mail

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(StationLetter, pk=letter_pk, schedule_block=block)

    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    if request.method == 'GET':
        return render(request, 'course/station_letter_loading.html', {
            'course': course, 'schedule_block': block, 'letter': letter,
        })

    template_obj = pick_active_letter_template(StationLetterTemplate, course.job_profile)
    if not template_obj:
        return JsonResponse({'error': 'Keine aktive Vorlage gefunden.'})

    assignments = (
        InternshipAssignment.objects
        .filter(schedule_block=block)
        .select_related('student__gender', 'student__address', 'unit', 'location', 'instructor')
        .order_by('student__last_name', 'student__first_name', 'start_date')
    )
    creator = letter.generated_by
    send_immediately = letter.status == BLOCK_LETTER_STATUS_SENT

    if not send_immediately:
        # Dokumente werden erst bei der Freigabe mit dem Namen der genehmigenden Person generiert
        detail_url = f'/kurs/{course.pk}/ablaufplan/{block.public_id}/stationsschreiben/{letter.pk}/'
        return JsonResponse({'redirect_url': detail_url})

    errors = []

    from document.contexts import student_context, course_context, block_context, einsatz_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    for assignment in assignments:
        student = assignment.student
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **einsatz_context(assignment),
                **creator_context(creator),
                **meta_context(),
                'anrede': (
                    f'{student.gender.description} {student.last_name},'
                    if student.gender else f'{student.first_name} {student.last_name},'
                ),
                'freitext': letter.free_text,
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Stationsschreiben: Rendering für %s/%s fehlgeschlagen: %s', student, assignment.unit, exc)
            errors.append(f'{student} – {assignment.unit}')
            continue

        title = f'Stationszuweisungsschreiben {assignment.unit.name} – {student.first_name} {student.last_name}'
        filename = (
            f'stationsschreiben_{student.last_name}_{student.first_name}'
            f'_{assignment.unit.name.replace(" ", "_")}.docx'
        )
        doc_id = upload_to_paperless(
            file_bytes=file_bytes, title=title, student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Stationsschreiben: Upload für %s/%s fehlgeschlagen', student, assignment.unit)
            errors.append(f'{student} – {assignment.unit}')
            continue

        item = StationLetterItem.objects.create(letter=letter, assignment=assignment, paperless_id=doc_id)

        if send_immediately and student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            attachments = [(f'{filename[:-5]}.pdf', pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Stationszuweisungsschreiben: {assignment.unit.name}',
                    body_text=(
                        f'Guten Tag {student.first_name} {student.last_name},\n\n'
                        f'im Anhang erhalten Sie Ihr Stationszuweisungsschreiben für den Einsatz '
                        f'bei {assignment.unit.name} ({assignment.start_date.strftime("%d.%m.%Y")} – '
                        f'{assignment.end_date.strftime("%d.%m.%Y")}).\n\n'
                        f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                    ),
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Stationsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    detail_url = f'/kurs/{course.pk}/ablaufplan/{block.public_id}/stationsschreiben/{letter.pk}/'
    if errors:
        return JsonResponse({'error': f'Fehler bei: {", ".join(errors)}', 'redirect_url': detail_url})
    return JsonResponse({'redirect_url': detail_url})


@login_required
def station_letter_detail(request, course_pk, block_public_id, letter_pk):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(StationLetter, pk=letter_pk, schedule_block=block)
    items = (
        letter.items
        .select_related('assignment__student__gender', 'assignment__unit')
        .order_by('assignment__student__last_name', 'assignment__student__first_name', 'assignment__start_date')
    )
    return render(request, 'course/station_letter_detail.html', {
        'course': course, 'schedule_block': block, 'letter': letter, 'items': items,
        'can_regenerate': _is_leitung_or_referat(request.user),
    })


@login_required
def station_letter_approve(request, course_pk, block_public_id, letter_pk):
    import logging
    from django.core.exceptions import PermissionDenied
    from django.utils import timezone
    from services.roles import is_training_director
    from services.paperless import PaperlessService
    from services.email import send_mail

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(StationLetter, pk=letter_pk, schedule_block=block)

    if not (request.user.is_staff or is_training_director(request.user)):
        raise PermissionDenied

    if request.method == 'GET':
        assignments = (
            InternshipAssignment.objects
            .filter(schedule_block=block)
            .select_related('student__gender', 'unit')
            .order_by('student__last_name', 'student__first_name', 'start_date')
        )
        return render(request, 'course/station_letter_approve.html', {
            'course': course, 'schedule_block': block, 'letter': letter, 'assignments': assignments,
        })

    # POST – Dokumente mit Unterzeichner generieren, in Paperless ablegen, versenden
    from document.contexts import student_context, course_context, block_context, einsatz_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    letter.approved_by = request.user
    letter.approved_at = timezone.now()
    letter.save(update_fields=['approved_by', 'approved_at'])

    template_obj = pick_active_letter_template(StationLetterTemplate, course.job_profile)
    if not template_obj:
        messages.error(request, 'Keine aktive Vorlage gefunden. Freigabe nicht möglich.')
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    assignments = (
        InternshipAssignment.objects
        .filter(schedule_block=block)
        .select_related('student__gender', 'student__address', 'unit', 'location', 'instructor')
        .order_by('student__last_name', 'student__first_name', 'start_date')
    )
    creator = letter.generated_by
    errors = []

    for assignment in assignments:
        student = assignment.student
        try:
            ctx = {
                **student_context(student),
                **course_context(course),
                **block_context(block),
                **einsatz_context(assignment),
                **creator_context(creator),
                **meta_context(),
                'anrede': (
                    f'{student.gender.description} {student.last_name},'
                    if student.gender else f'{student.first_name} {student.last_name},'
                ),
                'freitext': letter.free_text,
                'zeichnung': letter.approved_by.last_name,
            }
            file_bytes = render_docx(template_obj.template_file.path, ctx)
        except Exception as exc:
            logger.error('Freigabe Stationsschreiben: Rendering für %s/%s fehlgeschlagen: %s', student, assignment.unit, exc)
            errors.append(f'{student} – {assignment.unit}')
            continue

        title = f'Stationszuweisungsschreiben {assignment.unit.name} – {student.first_name} {student.last_name}'
        filename = (
            f'stationsschreiben_{student.last_name}_{student.first_name}'
            f'_{assignment.unit.name.replace(" ", "_")}.docx'
        )
        doc_id = upload_to_paperless(
            file_bytes=file_bytes, title=title, student_id=str(student.pk),
            filename=filename,
        )
        if doc_id is None:
            logger.error('Freigabe Stationsschreiben: Upload für %s/%s fehlgeschlagen', student, assignment.unit)
            errors.append(f'{student} – {assignment.unit}')
            continue

        item, _ = StationLetterItem.objects.get_or_create(letter=letter, assignment=assignment)
        if item.paperless_id and item.paperless_id != doc_id:
            PaperlessService.delete_document(item.paperless_id)
        item.paperless_id = doc_id
        item.email_sent = False
        item.save(update_fields=['paperless_id', 'email_sent'])

        if student.email_id:
            pdf_bytes = PaperlessService.download_pdf(doc_id)
            filename_pdf = (
                f'stationsschreiben_{student.last_name}_{student.first_name}'
                f'_{assignment.unit.name.replace(" ", "_")}.pdf'
            )
            attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
            try:
                send_mail(
                    subject=f'Stationszuweisungsschreiben: {assignment.unit.name}',
                    body_text=(
                        f'Guten Tag {student.first_name} {student.last_name},\n\n'
                        f'im Anhang erhalten Sie Ihr Stationszuweisungsschreiben für den Einsatz '
                        f'bei {assignment.unit.name} ({assignment.start_date.strftime("%d.%m.%Y")} – '
                        f'{assignment.end_date.strftime("%d.%m.%Y")}).\n\n'
                        f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                    ),
                    recipient_list=[student.email_id],
                    attachments=attachments or None,
                )
                item.email_sent = True
                item.save(update_fields=['email_sent'])
            except Exception as exc:
                logger.error('Freigabe Stationsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)
                errors.append(f'{student} – {assignment.unit}')

    letter.status = BLOCK_LETTER_STATUS_SENT
    letter.save(update_fields=['status'])

    if errors:
        messages.warning(request, f'Freigabe erteilt, aber Fehler bei: {", ".join(errors)}')
    else:
        messages.success(request, f'Stationszuweisungsschreiben für „{block.name}" freigegeben und versendet.')
    return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


# ── Einzelnes Schreiben neu erstellen ────────────────────────────────────────

def _regenerate_block_letter_item(letter, item, block, course, send_email):
    """Hilfsfunktion: Ein einzelnes Zuweisungsschreiben neu generieren und hochladen."""
    import logging
    from services.paperless import PaperlessService
    from services.email import send_mail
    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    logger = logging.getLogger(__name__)
    student = item.student

    template_obj = pick_active_letter_template(BlockLetterTemplate, course.job_profile)
    if not template_obj:
        return False, 'Keine aktive Vorlage gefunden.'

    creator = letter.generated_by

    try:
        ctx = {
            **student_context(student),
            **course_context(course),
            **block_context(block),
            **creator_context(creator),
            **meta_context(),
            'freitext': letter.free_text,
            'zeichnung': letter.approved_by.last_name,
        }
        file_bytes = render_docx(template_obj.template_file.path, ctx)
    except Exception as exc:
        logger.error('Neu erstellen Zuweisungsschreiben: Rendering für %s fehlgeschlagen: %s', student, exc)
        return False, f'Fehler beim Rendern: {exc}'

    title = f'Zuweisungsschreiben {block.name} – {student.first_name} {student.last_name}'
    filename = f'zuweisungsschreiben_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
    doc_id = upload_to_paperless(
        file_bytes=file_bytes, title=title, student_id=str(student.pk),
        filename=filename,
    )
    if doc_id is None:
        return False, 'Upload zu Paperless fehlgeschlagen.'

    if item.paperless_id and item.paperless_id != doc_id:
        PaperlessService.delete_document(item.paperless_id)
    item.paperless_id = doc_id
    item.email_sent = False
    item.save(update_fields=['paperless_id', 'email_sent'])

    if send_email and student.email_id:
        pdf_bytes = PaperlessService.download_pdf(doc_id)
        filename_pdf = f'zuweisungsschreiben_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.pdf'
        attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
        try:
            send_mail(
                subject=f'Zuweisungsschreiben: {block.name}',
                body_text=f'Guten Tag {student.first_name} {student.last_name},\n\nim Anhang erhalten Sie Ihr aktualisiertes Zuweisungsschreiben für den Block „{block.name}".\n\nMit freundlichen Grüßen\nIhr Azubi-Portal',
                recipient_list=[student.email_id],
                attachments=attachments or None,
            )
            item.email_sent = True
            item.save(update_fields=['email_sent'])
        except Exception as exc:
            logger.error('Neu erstellen Zuweisungsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    return True, None


@login_required
def block_letter_item_regenerate(request, course_pk, block_public_id, letter_pk, item_pk):
    """Einzelnes Zuweisungsschreiben neu erstellen (z.B. nach Datenänderung)."""
    from django.core.exceptions import PermissionDenied

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course)
    letter = get_object_or_404(BlockLetter, pk=letter_pk, schedule_block=block)
    item = get_object_or_404(BlockLetterItem, pk=item_pk, letter=letter)

    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    if request.method != 'POST':
        return redirect('course:block_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    if letter.status != BLOCK_LETTER_STATUS_SENT or not letter.approved_by:
        messages.error(request, 'Neu erstellen ist nur für freigegebene Schreiben möglich.')
        return redirect('course:block_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    send_email = request.POST.get('send_email') == '1'
    ok, err = _regenerate_block_letter_item(letter, item, block, course, send_email)
    if ok:
        suffix = ' und erneut versendet' if send_email else ''
        messages.success(request, f'Zuweisungsschreiben für {item.student} wurde neu erstellt{suffix}.')
    else:
        messages.error(request, f'Fehler: {err}')

    return redirect('course:block_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


@login_required
def internship_plan_item_regenerate(request, course_pk, block_public_id, letter_pk, item_pk):
    """Einzelnen Praktikumsplan neu erstellen (z.B. nach Einsatzänderung)."""
    import logging
    from django.core.exceptions import PermissionDenied
    from services.paperless import PaperlessService
    from services.email import send_mail
    from document.contexts import student_context, course_context, block_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(InternshipPlanLetter, pk=letter_pk, schedule_block=block)
    item = get_object_or_404(InternshipPlanItem, pk=item_pk, letter=letter)

    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    if request.method != 'POST':
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    if letter.status != BLOCK_LETTER_STATUS_SENT or not letter.approved_by:
        messages.error(request, 'Neu erstellen ist nur für freigegebene Schreiben möglich.')
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    send_email = request.POST.get('send_email') == '1'
    student = item.student

    template_obj = pick_active_letter_template(InternshipPlanTemplate, course.job_profile)
    if not template_obj:
        messages.error(request, 'Keine aktive Vorlage gefunden.')
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    assignments = (
        InternshipAssignment.objects
        .filter(schedule_block=block, student=student)
        .select_related('unit', 'location', 'instructor')
        .order_by('start_date')
    )
    einsaetze = [
        {
            'einheit': a.unit.name,
            'beginn': a.start_date.strftime('%d.%m.%Y'),
            'ende': a.end_date.strftime('%d.%m.%Y'),
            'standort': str(a.location) if a.location else '',
            'praxistutor': f'{a.instructor.first_name} {a.instructor.last_name}' if a.instructor else '',
        }
        for a in assignments
    ]

    creator = letter.generated_by
    try:
        ctx = {
            **student_context(student),
            **course_context(course),
            **block_context(block),
            **creator_context(creator),
            **meta_context(),
            'anrede': (
                f'{student.gender.description} {student.last_name},'
                if student.gender else f'{student.first_name} {student.last_name},'
            ),
            'freitext': letter.free_text,
            'einsaetze': einsaetze,
            'zeichnung': letter.approved_by.last_name,
        }
        file_bytes = render_docx(template_obj.template_file.path, ctx)
    except Exception as exc:
        logger.error('Neu erstellen Praktikumsplan: Rendering für %s fehlgeschlagen: %s', student, exc)
        messages.error(request, f'Fehler beim Rendern: {exc}')
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    title = f'Praktikumsplan {block.name} – {student.first_name} {student.last_name}'
    filename = f'praktikumsplan_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.docx'
    doc_id = upload_to_paperless(
        file_bytes=file_bytes, title=title, student_id=str(student.pk),
        filename=filename,
    )
    if doc_id is None:
        messages.error(request, 'Upload zu Paperless fehlgeschlagen.')
        return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    if item.paperless_id and item.paperless_id != doc_id:
        PaperlessService.delete_document(item.paperless_id)
    item.paperless_id = doc_id
    item.email_sent = False
    item.save(update_fields=['paperless_id', 'email_sent'])

    if send_email and student.email_id:
        pdf_bytes = PaperlessService.download_pdf(doc_id)
        filename_pdf = f'praktikumsplan_{block.name.replace(" ", "_")}_{student.last_name}_{student.first_name}.pdf'
        attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
        try:
            send_mail(
                subject=f'Praktikumsplan: {block.name}',
                body_text=(
                    f'Guten Tag {student.first_name} {student.last_name},\n\n'
                    f'im Anhang erhalten Sie Ihren aktualisierten Praktikumsplan für den Block „{block.name}".\n\n'
                    f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                ),
                recipient_list=[student.email_id],
                attachments=attachments or None,
            )
            item.email_sent = True
            item.save(update_fields=['email_sent'])
        except Exception as exc:
            logger.error('Neu erstellen Praktikumsplan: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    suffix = ' und erneut versendet' if send_email else ''
    messages.success(request, f'Praktikumsplan für {student} wurde neu erstellt{suffix}.')
    return redirect('course:internship_plan_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


@login_required
def station_letter_item_regenerate(request, course_pk, block_public_id, letter_pk, item_pk):
    """Einzelnes Stationsschreiben neu erstellen (z.B. nach Einsatzänderung)."""
    import logging
    from django.core.exceptions import PermissionDenied
    from services.paperless import PaperlessService
    from services.email import send_mail
    from document.contexts import student_context, course_context, block_context, einsatz_context, creator_context, meta_context
    from document.render import render_docx, upload_to_paperless

    logger = logging.getLogger(__name__)
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    letter = get_object_or_404(StationLetter, pk=letter_pk, schedule_block=block)
    item = get_object_or_404(StationLetterItem, pk=item_pk, letter=letter)

    if not _is_leitung_or_referat(request.user):
        raise PermissionDenied

    if request.method != 'POST':
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    if letter.status != BLOCK_LETTER_STATUS_SENT or not letter.approved_by:
        messages.error(request, 'Neu erstellen ist nur für freigegebene Schreiben möglich.')
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    send_email = request.POST.get('send_email') == '1'
    assignment = item.assignment
    student = assignment.student

    template_obj = pick_active_letter_template(StationLetterTemplate, course.job_profile)
    if not template_obj:
        messages.error(request, 'Keine aktive Vorlage gefunden.')
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    creator = letter.generated_by
    try:
        ctx = {
            **student_context(student),
            **course_context(course),
            **block_context(block),
            **einsatz_context(assignment),
            **creator_context(creator),
            **meta_context(),
            'anrede': (
                f'{student.gender.description} {student.last_name},'
                if student.gender else f'{student.first_name} {student.last_name},'
            ),
            'freitext': letter.free_text,
            'zeichnung': letter.approved_by.last_name,
        }
        file_bytes = render_docx(template_obj.template_file.path, ctx)
    except Exception as exc:
        logger.error('Neu erstellen Stationsschreiben: Rendering für %s/%s fehlgeschlagen: %s', student, assignment.unit, exc)
        messages.error(request, f'Fehler beim Rendern: {exc}')
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    title = f'Stationszuweisungsschreiben {assignment.unit.name} – {student.first_name} {student.last_name}'
    filename = (
        f'stationsschreiben_{student.last_name}_{student.first_name}'
        f'_{assignment.unit.name.replace(" ", "_")}.docx'
    )
    doc_id = upload_to_paperless(
        file_bytes=file_bytes, title=title, student_id=str(student.pk),
        filename=filename,
    )
    if doc_id is None:
        messages.error(request, 'Upload zu Paperless fehlgeschlagen.')
        return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)

    if item.paperless_id and item.paperless_id != doc_id:
        PaperlessService.delete_document(item.paperless_id)
    item.paperless_id = doc_id
    item.email_sent = False
    item.save(update_fields=['paperless_id', 'email_sent'])

    if send_email and student.email_id:
        pdf_bytes = PaperlessService.download_pdf(doc_id)
        filename_pdf = (
            f'stationsschreiben_{student.last_name}_{student.first_name}'
            f'_{assignment.unit.name.replace(" ", "_")}.pdf'
        )
        attachments = [(filename_pdf, pdf_bytes, 'application/pdf')] if pdf_bytes else []
        try:
            send_mail(
                subject=f'Stationszuweisungsschreiben: {assignment.unit.name}',
                body_text=(
                    f'Guten Tag {student.first_name} {student.last_name},\n\n'
                    f'im Anhang erhalten Sie Ihr aktualisiertes Stationszuweisungsschreiben für den Einsatz '
                    f'bei {assignment.unit.name} ({assignment.start_date.strftime("%d.%m.%Y")} – '
                    f'{assignment.end_date.strftime("%d.%m.%Y")}).\n\n'
                    f'Mit freundlichen Grüßen\nIhr Azubi-Portal'
                ),
                recipient_list=[student.email_id],
                attachments=attachments or None,
            )
            item.email_sent = True
            item.save(update_fields=['email_sent'])
        except Exception as exc:
            logger.error('Neu erstellen Stationsschreiben: E-Mail an %s fehlgeschlagen: %s', student.email_id, exc)

    suffix = ' und erneut versendet' if send_email else ''
    messages.success(request, f'Stationsschreiben für {student} ({assignment.unit}) wurde neu erstellt{suffix}.')
    return redirect('course:station_letter_detail', course_pk=course.pk, block_public_id=block.public_id, letter_pk=letter.pk)


# ── Kapazitätsplanung ────────────────────────────────────────────────────────

@login_required
def capacity_planning(request):
    from calendar import monthrange
    from services.roles import (
        is_training_director, is_training_office, is_training_coordinator,
        get_chief_instructor,
    )
    from organisation.models import OrganisationalUnit

    is_leitung = is_training_director(request.user)
    is_referat = is_training_office(request.user)
    is_koord = is_training_coordinator(request.user)

    if not (is_leitung or is_referat or is_koord):
        raise PermissionDenied

    koord_unit_pks = None
    if is_koord and not (is_leitung or is_referat):
        from instructor.views import _get_coordination_area
        chief = get_chief_instructor(request.user)
        if chief and chief.coordination:
            descendant_pks, _, _ = _get_coordination_area(chief.coordination)
            koord_unit_pks = descendant_pks
        else:
            koord_unit_pks = []

    today = date.today()

    # Startmonat aus URL-Parameter oder aktueller Monat
    month_param = request.GET.get('month', '')
    try:
        start_year, start_month = map(int, month_param.split('-'))
        start = date(start_year, start_month, 1)
    except (ValueError, AttributeError):
        start = date(today.year, today.month, 1)

    # 12 aufeinanderfolgende Monate aufbauen
    months = []
    y, m = start.year, start.month
    for _ in range(12):
        last_day = monthrange(y, m)[1]
        months.append((date(y, m, 1), date(y, m, last_day)))
        m += 1
        if m > 12:
            m = 1
            y += 1

    period_start = months[0][0]
    period_end = months[-1][1]

    # Alle Einsätze in Praktikumsblöcken, die in den Zeitraum fallen
    assignments_qs = InternshipAssignment.objects.filter(
        schedule_block__block_type='internship',
        start_date__lte=period_end,
        end_date__gte=period_start,
    ).select_related('unit')
    if koord_unit_pks is not None:
        assignments_qs = assignments_qs.filter(unit_id__in=koord_unit_pks)

    # Zeitspannen pro Einheit vorhalten
    unit_ranges: dict[int, list[tuple]] = {}
    for a in assignments_qs:
        unit_ranges.setdefault(a.unit_id, []).append((a.start_date, a.end_date))

    units_qs = OrganisationalUnit.objects.all()
    if koord_unit_pks is not None:
        units_qs = units_qs.filter(pk__in=koord_unit_pks)

    rows = []
    for unit in units_qs.order_by('name'):
        ranges = unit_ranges.get(unit.pk, [])
        if not ranges and unit.max_capacity is None:
            continue  # Einheiten ohne Einsätze und ohne Kapazitätsangabe ausblenden

        cells = []
        for month_start, month_end in months:
            count = sum(
                1 for s, e in ranges
                if s <= month_end and e >= month_start
            )
            cap = unit.max_capacity
            if cap:
                if count == 0:
                    status = 'free'
                elif count >= cap:
                    status = 'danger'
                elif count / cap >= 0.75:
                    status = 'warning'
                else:
                    status = 'success'
            else:
                status = 'has_count' if count > 0 else 'free'
            cells.append({
                'count': count,
                'cap': cap,
                'status': status,
                'is_today': month_start.year == today.year and month_start.month == today.month,
            })
        STATUS_RANK = {'danger': 3, 'warning': 2, 'success': 1, 'has_count': 1, 'free': 0}
        worst = max((c['status'] for c in cells), key=lambda s: STATUS_RANK.get(s, 0))
        has_any = any(c['count'] > 0 for c in cells)
        rows.append({'unit': unit, 'cells': cells, 'worst': worst, 'has_any': has_any})

    # Konflikt-Übersicht: Einheiten mit Überkapazität in den nächsten 365 Tagen
    conflict_summary = []
    conflict_map: dict[int, int] = {}
    for row in rows:
        unit = row['unit']
        if not unit.max_capacity:
            continue
        ranges = unit_ranges.get(unit.pk, [])
        if not ranges:
            continue
        days_over = 0
        next_conflict = None
        for i in range(366):
            d = today + timedelta(days=i)
            count = sum(1 for s, e in ranges if s <= d <= e)
            if count >= unit.max_capacity:
                days_over += 1
                if next_conflict is None:
                    next_conflict = d
        if days_over > 0:
            conflict_map[unit.pk] = days_over
            conflict_summary.append({
                'unit': unit,
                'days_over': days_over,
                'next_conflict': next_conflict,
            })
    conflict_summary.sort(key=lambda x: -x['days_over'])
    for row in rows:
        row['days_over'] = conflict_map.get(row['unit'].pk, 0)

    # Navigation: jeweils einen Monat vor/zurück
    def shift_month(d, delta):
        total = d.month + delta
        y = d.year + (total - 1) // 12
        m = (total - 1) % 12 + 1
        return date(y, m, 1)

    prev_start = shift_month(start, -1)
    next_start = shift_month(start, 1)

    return render(request, 'course/capacity_planning.html', {
        'months': [m[0] for m in months],
        'rows': rows,
        'start': start,
        'today': today,
        'prev_month': prev_start.strftime('%Y-%m'),
        'next_month': next_start.strftime('%Y-%m'),
        'conflict_summary': conflict_summary,
    })


# ── Kurskalender ─────────────────────────────────────────────────────────────

@login_required
def course_calendar(request):
    from datetime import date, timedelta
    from services.roles import (
        is_training_director, is_training_office, is_training_coordinator,
        get_chief_instructor,
    )

    is_leitung = is_training_director(request.user)
    is_referat = is_training_office(request.user)
    is_koord = is_training_coordinator(request.user)

    if not (is_leitung or is_referat or is_koord):
        raise PermissionDenied

    today = date.today()
    year = int(request.GET.get('year', today.year))
    year_start = date(year, 1, 1)
    year_end = date(year, 12, 31)
    year_days = (year_end - year_start).days + 1

    # Monatskopfzeile
    months_data = []
    for m in range(1, 13):
        m_start = date(year, m, 1)
        m_end = date(year, m + 1, 1) - timedelta(days=1) if m < 12 else year_end
        m_days = (m_end - m_start).days + 1
        months_data.append({
            'name': m_start.strftime('%B'),
            'short': m_start.strftime('%b'),
            'width': f'{m_days / year_days * 100:.4f}',
            'offset': f'{(m_start - year_start).days / year_days * 100:.4f}',
        })

    today_offset = None
    if year_start <= today <= year_end:
        today_offset = f'{(today - year_start).days / year_days * 100:.4f}'

    # Kurse und Blöcke laden
    courses_qs = Course.objects.filter(
        end_date__gte=year_start,
        start_date__lte=year_end,
    ).prefetch_related('schedule_blocks').order_by('start_date')

    # Koordination: nur Kurse mit Einsätzen in ihrer Einheit
    if is_koord and not (is_leitung or is_referat):
        from instructor.views import _get_coordination_area
        chief = get_chief_instructor(request.user)
        if chief and chief.coordination:
            descendant_pks, _, _ = _get_coordination_area(chief.coordination)
            block_pks = (
                InternshipAssignment.objects
                .filter(unit_id__in=descendant_pks)
                .values_list('schedule_block_id', flat=True)
                .distinct()
            )
            course_pks = (
                ScheduleBlock.objects
                .filter(pk__in=block_pks)
                .values_list('course_id', flat=True)
                .distinct()
            )
            courses_qs = courses_qs.filter(pk__in=course_pks)
        else:
            courses_qs = Course.objects.none()

    from services.colors import BUNDESFARBEN_BY_NAME as _BF
    color_lesson = _BF['Petrol']
    color_internship = _BF['Dunkelrot']
    color_today = _BF['Rot']

    calendar_rows = []
    bar_height = 32
    bar_gap = 6

    for course in courses_qs:
        blocks = course.schedule_blocks.filter(
            end_date__gte=year_start,
            start_date__lte=year_end,
        ).order_by('start_date')

        bars = []
        for block in blocks:
            clip_start = max(block.start_date, year_start)
            clip_end = min(block.end_date, year_end)
            if clip_start > clip_end:
                continue
            raw_offset = (clip_start - year_start).days / year_days * 100
            raw_width = max(((clip_end - clip_start).days + 1) / year_days * 100, 0.5)
            color = color_internship if block.is_internship else color_lesson
            bars.append({
                'block': block,
                'offset': f'{raw_offset:.4f}',
                'width': f'{raw_width:.4f}',
                'color': color,
                'label': block.name,
                'start_fmt': block.start_date.strftime('%d.%m.%Y'),
                'end_fmt': block.end_date.strftime('%d.%m.%Y'),
            })

        row_height = bar_height + 2 * bar_gap if bars else bar_height
        calendar_rows.append({
            'course': course,
            'bars': bars,
            'row_height': row_height,
        })

    return render(request, 'course/course_calendar.html', {
        'year': year,
        'prev_year': year - 1,
        'next_year': year + 1,
        'months_data': months_data,
        'today_offset': today_offset,
        'calendar_rows': calendar_rows,
        'bar_height': bar_height,
        'bar_gap': bar_gap,
        'color_lesson': color_lesson,
        'color_internship': color_internship,
        'color_today': color_today,
    })


# ── Checklisten auf Kurs-Ebene ────────────────────────────────────────────────

@login_required
@require_POST
def course_checklist_bulk_create(request, pk):
    from services.roles import is_training_director, is_training_office
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied
    course = get_object_or_404(Course, pk=pk)
    from student.models import ChecklistTemplate, StudentChecklist, StudentChecklistItem

    template_pk = request.POST.get('template')
    template = get_object_or_404(ChecklistTemplate, pk=template_pk, is_active=True)
    students = Student.objects.filter(course=course)

    created = 0
    skipped = 0
    for student in students:
        already_exists = StudentChecklist.objects.filter(
            student=student, template=template
        ).exists()
        if already_exists:
            skipped += 1
            continue
        checklist = StudentChecklist.objects.create(
            student=student,
            template=template,
            name=template.name,
            created_by=request.user,
        )
        for item in template.items.order_by('order', 'text'):
            StudentChecklistItem.objects.create(
                checklist=checklist,
                text=item.text,
                order=item.order,
            )
        created += 1

    if created:
        messages.success(
            request,
            f'Checkliste „{template.name}" für {created} Nachwuchskraft/kräfte angelegt.'
            + (f' ({skipped} bereits vorhanden, übersprungen.)' if skipped else '')
        )
    else:
        messages.warning(request, f'Alle {skipped} Nachwuchskräfte haben diese Checkliste bereits.')

    return redirect(f'/kurs/{course.pk}/?tab=checklisten')


# ── Kurs-eigene Checklisten ───────────────────────────────────────────────────

@login_required
@require_POST
def course_checklist_create(request, pk):
    """Legt eine kurseigene Checkliste aus einer Vorlage an."""
    from services.roles import is_training_director, is_training_office
    from student.models import ChecklistTemplate
    from course.models import CourseChecklist, CourseChecklistItem
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied
    course = get_object_or_404(Course, pk=pk)
    template_pk = request.POST.get('template')
    template = get_object_or_404(ChecklistTemplate, pk=template_pk, is_active=True)

    already = CourseChecklist.objects.filter(course=course, template=template).exists()
    if already:
        messages.warning(request, f'Die Checkliste „{template.name}" existiert für diesen Kurs bereits.')
        return redirect(f'/kurs/{course.pk}/?tab=kurs-checklisten')

    checklist = CourseChecklist.objects.create(
        course=course,
        template=template,
        name=template.name,
        created_by=request.user,
    )
    for item in template.items.order_by('order', 'text'):
        CourseChecklistItem.objects.create(
            checklist=checklist,
            text=item.text,
            order=item.order,
        )
    messages.success(request, f'Kurs-Checkliste „{template.name}" wurde angelegt.')
    return redirect(f'/kurs/{course.pk}/?tab=kurs-checklisten')


@login_required
@require_POST
def course_checklist_item_toggle(request, pk, checklist_pk, item_pk):
    """Setzt einen Kurs-Checklisten-Punkt auf erledigt/nicht erledigt."""
    from django.utils import timezone
    from services.roles import is_training_director, is_training_office
    from course.models import CourseChecklist, CourseChecklistItem
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied
    course = get_object_or_404(Course, pk=pk)
    checklist = get_object_or_404(CourseChecklist, pk=checklist_pk, course=course)
    item = get_object_or_404(CourseChecklistItem, pk=item_pk, checklist=checklist)

    if item.completed:
        item.completed = False
        item.completed_at = None
        item.completed_by = None
    else:
        item.completed = True
        item.completed_at = timezone.now()
        item.completed_by = request.user
    item.save()
    return redirect(f'/kurs/{course.pk}/?tab=kurs-checklisten')


@login_required
@require_POST
def course_checklist_delete(request, pk, checklist_pk):
    """Löscht eine kurseigene Checkliste."""
    from services.roles import is_training_director, is_training_office
    from course.models import CourseChecklist
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied
    course = get_object_or_404(Course, pk=pk)
    checklist = get_object_or_404(CourseChecklist, pk=checklist_pk, course=course)
    name = checklist.name
    checklist.delete()
    messages.success(request, f'Kurs-Checkliste „{name}" wurde gelöscht.')
    return redirect(f'/kurs/{course.pk}/?tab=kurs-checklisten')


# ── Ausbildungsplan (Soll-Ist-Abgleich) ─────────────────────────────────────

@login_required
def curriculum_overview(request, pk):
    """Soll-Ist-Abgleich des Ausbildungsplans für einen Kurs."""
    from services.roles import is_training_director, is_training_office
    if not (is_training_director(request.user) or is_training_office(request.user)):
        raise PermissionDenied

    course = get_object_or_404(Course.objects.select_related('job_profile'), pk=pk)
    from .curriculum import get_course_curriculum_overview
    overview = get_course_curriculum_overview(course)

    return render(request, 'course/curriculum_overview.html', {
        'course': course,
        'requirements': overview['requirements'],
        'student_rows': overview['students'],
    })


# ── Einsatzvorschläge ────────────────────────────────────────────────────────

@login_required
def internship_suggestions(request, course_pk, block_public_id):
    """Vorschläge für Praktikumseinsätze generieren."""
    from services.roles import is_training_director, is_training_office
    if not (is_training_director(request.user) or is_training_office(request.user)):
        from django.core.exceptions import PermissionDenied
        raise PermissionDenied

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(ScheduleBlock, public_id=block_public_id, course=course, block_type='internship')
    student_pk = request.GET.get('student', '')

    from student.models import Student
    students = Student.objects.filter(course=course, anonymized_at__isnull=True).order_by('last_name', 'first_name')
    selected_student = None
    suggestions = []

    if student_pk:
        selected_student = students.filter(pk=student_pk).first()

    if selected_student:
        full_pks, _ = _get_unit_capacity_info(block, start_date=block.start_date, end_date=block.end_date)
        from .suggestions import generate_suggestions
        suggestions = generate_suggestions(selected_student, block, full_unit_pks=full_pks)

    return render(request, 'course/internship_suggestions.html', {
        'course': course,
        'schedule_block': block,
        'students': students,
        'selected_student': selected_student,
        'suggestions': suggestions,
    })




# ── Berufsbild-Konfiguration (Kriterien + Kompetenz-Endziele, Frontend) ─────

def _job_profile_admin_required(request):
    """Schutz: nur Ausbildungsleitung/-referat darf Berufsbild-Konfiguration ändern."""
    from django.core.exceptions import PermissionDenied
    from services.roles import is_training_director, is_training_office
    if not (request.user.is_authenticated and (
        request.user.is_staff
        or is_training_director(request.user)
        or is_training_office(request.user)
    )):
        raise PermissionDenied


@login_required
def job_profile_list(request):
    """Übersicht aller Berufsbilder mit Pflege-Links."""
    _job_profile_admin_required(request)
    from course.models import JobProfile, CompetenceTarget
    from assessment.models import AssessmentCriterion
    from django.db.models import Count
    profiles = JobProfile.objects.annotate(
        criterion_count=Count('assessment_criteria', distinct=True),
        target_count=Count('competence_targets', distinct=True),
    ).order_by('description')
    return render(request, 'course/job_profile_list.html', {'profiles': profiles})


@login_required
def job_profile_config(request, pk):
    """Konfigurations-Übersicht eines Berufsbilds: Kriterien + Endziele."""
    _job_profile_admin_required(request)
    from course.models import JobProfile, CompetenceTarget
    from assessment.models import AssessmentCriterion

    profile = get_object_or_404(JobProfile, pk=pk)
    criteria = (
        AssessmentCriterion.objects
        .filter(job_profile=profile)
        .prefetch_related('competence_weights__competence')
        .order_by('category', 'order', 'name')
    )
    targets = (
        CompetenceTarget.objects
        .filter(job_profile=profile)
        .select_related('competence')
        .order_by('competence__name')
    )
    return render(request, 'course/job_profile_config.html', {
        'profile':  profile,
        'criteria': criteria,
        'targets':  targets,
    })


def _criterion_weight_formset(criterion=None):
    """Formset für CriterionCompetenceWeight (Mapping Kompetenz + Gewicht)."""
    from django.forms import inlineformset_factory
    from assessment.models import AssessmentCriterion, CriterionCompetenceWeight
    return inlineformset_factory(
        AssessmentCriterion,
        CriterionCompetenceWeight,
        fields=('competence', 'weight'),
        extra=2,
        can_delete=True,
    )


@login_required
def criterion_create(request, profile_pk):
    """Neues Beurteilungskriterium für ein Berufsbild + Kompetenz-Mapping."""
    _job_profile_admin_required(request)
    from course.models import JobProfile
    from assessment.models import AssessmentCriterion
    from assessment.forms import AssessmentCriterionForm

    profile = get_object_or_404(JobProfile, pk=profile_pk)
    FormSet = _criterion_weight_formset()

    if request.method == 'POST':
        form = AssessmentCriterionForm(request.POST)
        if form.is_valid():
            crit = form.save(commit=False)
            crit.job_profile = profile
            crit.save()
            formset = FormSet(request.POST, instance=crit)
            if formset.is_valid():
                formset.save()
                messages.success(request, f'Kriterium „{crit.name}" wurde angelegt.')
                return redirect('course:job_profile_config', pk=profile.pk)
            crit.delete()  # Rollback bei FormSet-Fehler
        formset = FormSet(request.POST)
    else:
        form = AssessmentCriterionForm()
        formset = FormSet()

    return render(request, 'course/criterion_form.html', {
        'profile': profile, 'form': form, 'formset': formset, 'action': 'Anlegen',
    })


@login_required
def criterion_edit(request, profile_pk, pk):
    _job_profile_admin_required(request)
    from course.models import JobProfile
    from assessment.models import AssessmentCriterion
    from assessment.forms import AssessmentCriterionForm

    profile = get_object_or_404(JobProfile, pk=profile_pk)
    crit = get_object_or_404(AssessmentCriterion, pk=pk, job_profile=profile)
    FormSet = _criterion_weight_formset()

    if request.method == 'POST':
        form = AssessmentCriterionForm(request.POST, instance=crit)
        formset = FormSet(request.POST, instance=crit)
        if form.is_valid() and formset.is_valid():
            form.save()
            formset.save()
            messages.success(request, f'Kriterium „{crit.name}" wurde aktualisiert.')
            return redirect('course:job_profile_config', pk=profile.pk)
    else:
        form = AssessmentCriterionForm(instance=crit)
        formset = FormSet(instance=crit)

    return render(request, 'course/criterion_form.html', {
        'profile': profile, 'form': form, 'formset': formset,
        'action': 'Bearbeiten', 'criterion': crit,
    })


@login_required
def criterion_delete(request, profile_pk, pk):
    _job_profile_admin_required(request)
    from course.models import JobProfile
    from assessment.models import AssessmentCriterion

    profile = get_object_or_404(JobProfile, pk=profile_pk)
    crit = get_object_or_404(AssessmentCriterion, pk=pk, job_profile=profile)
    if request.method == 'POST':
        name = crit.name
        try:
            crit.delete()
            messages.success(request, f'Kriterium „{name}" wurde gelöscht.')
        except Exception as exc:
            messages.error(request, f'Kriterium konnte nicht gelöscht werden: {exc}')
        return redirect('course:job_profile_config', pk=profile.pk)
    return render(request, 'course/criterion_confirm_delete.html', {
        'profile': profile, 'criterion': crit,
    })


@login_required
def competence_target_create(request, profile_pk):
    _job_profile_admin_required(request)
    from course.models import JobProfile
    from assessment.forms import CompetenceTargetForm

    profile = get_object_or_404(JobProfile, pk=profile_pk)
    if request.method == 'POST':
        form = CompetenceTargetForm(request.POST, job_profile=profile)
        if form.is_valid():
            target = form.save(commit=False)
            target.job_profile = profile
            target.save()
            messages.success(request, f'Endziel für „{target.competence.name}" wurde angelegt.')
            return redirect('course:job_profile_config', pk=profile.pk)
    else:
        form = CompetenceTargetForm(job_profile=profile)
    return render(request, 'course/target_form.html', {
        'profile': profile, 'form': form, 'action': 'Anlegen',
    })


@login_required
def competence_target_edit(request, profile_pk, pk):
    _job_profile_admin_required(request)
    from course.models import JobProfile, CompetenceTarget
    from assessment.forms import CompetenceTargetForm

    profile = get_object_or_404(JobProfile, pk=profile_pk)
    target = get_object_or_404(CompetenceTarget, pk=pk, job_profile=profile)
    if request.method == 'POST':
        form = CompetenceTargetForm(request.POST, instance=target, job_profile=profile)
        if form.is_valid():
            form.save()
            messages.success(request, f'Endziel für „{target.competence.name}" wurde aktualisiert.')
            return redirect('course:job_profile_config', pk=profile.pk)
    else:
        form = CompetenceTargetForm(instance=target, job_profile=profile)
    return render(request, 'course/target_form.html', {
        'profile': profile, 'form': form, 'action': 'Bearbeiten', 'target': target,
    })


@login_required
def competence_target_delete(request, profile_pk, pk):
    _job_profile_admin_required(request)
    from course.models import JobProfile, CompetenceTarget
    profile = get_object_or_404(JobProfile, pk=profile_pk)
    target = get_object_or_404(CompetenceTarget, pk=pk, job_profile=profile)
    if request.method == 'POST':
        target.delete()
        messages.success(request, 'Endziel wurde entfernt.')
        return redirect('course:job_profile_config', pk=profile.pk)
    return render(request, 'course/target_confirm_delete.html', {
        'profile': profile, 'target': target,
    })


# ── Seminar / Vortragsplanung ───────────────────────────────────────────────

def _annotate_lecture_layout(lectures, grid_start_hour, hour_height_px=60):
    """Setzt ``top_offset_px``, ``height_px`` sowie ``local_date`` (lokaler
    Kalendertag) und ``local_start``/``local_end`` (lokale datetimes) auf jedem
    Vortrag in-place. Pixel-Berechnung in lokaler Zeit – die DB liefert UTC.
    """
    from django.utils import timezone
    for lec in lectures:
        start_local = timezone.localtime(lec.start_datetime)
        end_local = timezone.localtime(lec.end_datetime)
        lec.local_start = start_local
        lec.local_end = end_local
        lec.local_date = start_local.date()
        start_minutes_in_grid = (
            (start_local.hour - grid_start_hour) * 60 + start_local.minute
        )
        duration_minutes = int((end_local - start_local).total_seconds() // 60)
        lec.top_offset_px = max(0, start_minutes_in_grid * hour_height_px // 60)
        lec.height_px = max(20, duration_minutes * hour_height_px // 60)


def _build_seminar_weeks(block, lectures):
    """Gruppiert die Wochentage des Blocks (Mo–Fr) in Kalenderwochen und ordnet
    jedem Tag die zugehörigen Vorträge zu (basierend auf dem lokalen Kalendertag).
    Setzt ``_annotate_lecture_layout`` voraus (für ``local_date``).
    """
    from collections import defaultdict
    by_date = defaultdict(list)
    for lec in lectures:
        by_date[lec.local_date].append(lec)

    weeks = []
    cur = block.start_date
    end = block.end_date
    week_start = cur - timedelta(days=cur.weekday())
    while week_start <= end:
        days = []
        for offset in range(5):  # Mo-Fr
            day = week_start + timedelta(days=offset)
            in_block = block.start_date <= day <= end
            days.append({
                'date': day,
                'in_block': in_block,
                'lectures': sorted(by_date.get(day, []), key=lambda l: l.local_start),
            })
        weeks.append({
            'iso_week': week_start.isocalendar().week,
            'iso_year': week_start.isocalendar().year,
            'monday': week_start,
            'friday': week_start + timedelta(days=4),
            'days': days,
        })
        week_start += timedelta(days=7)
    return weeks


def _seminar_grid_bounds(lectures, default_start_hour=8, default_end_hour=18):
    """Bestimmt die kleinste/größte Stunde, die im Stundenraster angezeigt wird –
    in lokaler Zeit."""
    from django.utils import timezone
    if not lectures:
        return default_start_hour, default_end_hour
    min_h = min(timezone.localtime(l.start_datetime).hour for l in lectures)
    max_h_values = []
    for l in lectures:
        end_local = timezone.localtime(l.end_datetime)
        max_h_values.append(end_local.hour + (1 if end_local.minute else 0))
    return min(default_start_hour, min_h), max(default_end_hour, max(max_h_values))


@login_required
def seminar_calendar(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(
        ScheduleBlock, public_id=block_public_id, course=course, block_type='seminar'
    )
    lectures = list(block.lectures.all().order_by('start_datetime'))
    start_hour, end_hour = _seminar_grid_bounds(lectures)
    _annotate_lecture_layout(lectures, start_hour)
    weeks = _build_seminar_weeks(block, lectures)
    hours = list(range(start_hour, end_hour + 1))
    return render(request, 'course/seminar_calendar.html', {
        'course': course,
        'schedule_block': block,
        'weeks': weeks,
        'lectures': lectures,
        'hours': hours,
        'grid_start_hour': start_hour,
        'grid_end_hour': end_hour,
    })


@login_required
def lecture_create(request, course_pk, block_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(
        ScheduleBlock, public_id=block_public_id, course=course, block_type='seminar'
    )
    initial = {}
    prefill_date = request.GET.get('date')
    if prefill_date:
        initial['lecture_date'] = prefill_date
    prefill_time = request.GET.get('time')
    if prefill_time:
        initial['start_time'] = prefill_time
    form = SeminarLectureForm(request.POST or None, block=block, initial=initial)
    if form.is_valid():
        lecture = form.save(commit=False)
        lecture.created_by = request.user
        lecture.save()
        from services.notifications import notify_lecture_request
        notify_lecture_request(request, lecture)
        messages.success(
            request,
            f'Vortrag „{lecture.topic}" wurde angelegt. Bestätigungs-E-Mail wurde an '
            f'{lecture.speaker_email} gesendet.'
        )
        return redirect('course:seminar_calendar', course_pk=course.pk,
                        block_public_id=block.public_id)
    return render(request, 'course/lecture_form.html', {
        'form': form,
        'course': course,
        'schedule_block': block,
        'action': 'Anlegen',
    })


@login_required
def lecture_edit(request, course_pk, block_public_id, lecture_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(
        ScheduleBlock, public_id=block_public_id, course=course, block_type='seminar'
    )
    lecture = get_object_or_404(SeminarLecture, public_id=lecture_public_id, schedule_block=block)
    old_start = lecture.start_datetime
    old_end = lecture.end_datetime
    form = SeminarLectureForm(request.POST or None, instance=lecture, block=block)
    if form.is_valid():
        new_lecture = form.save(commit=False)
        relevant_changed = (
            new_lecture.start_datetime != old_start
            or new_lecture.end_datetime != old_end
            or new_lecture.location != lecture.location
            or new_lecture.topic != lecture.topic
        )
        if relevant_changed and lecture.status != LECTURE_STATUS_DECLINED:
            new_lecture.notification_sequence = lecture.notification_sequence + 1
        new_lecture.save()
        if relevant_changed and lecture.status != LECTURE_STATUS_DECLINED:
            from services.notifications import notify_lecture_update
            notify_lecture_update(request, new_lecture)
            messages.info(request, 'Aktualisierte Termin-Einladung wurde an den Vortragenden gesendet.')
        messages.success(request, f'Vortrag „{new_lecture.topic}" wurde gespeichert.')
        return redirect('course:seminar_calendar', course_pk=course.pk,
                        block_public_id=block.public_id)
    return render(request, 'course/lecture_form.html', {
        'form': form,
        'course': course,
        'schedule_block': block,
        'lecture': lecture,
        'action': 'Bearbeiten',
    })


def _docx_set_cell_shading(cell, hex_color):
    """Setzt die Hintergrundfarbe einer Tabellenzelle (z.B. '1F4E79')."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _docx_set_table_borders(table, hex_color='2C3E50', size_eighths='8'):
    """Setzt einheitliche Rahmenlinien rund um alle Zellen einer Tabelle."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl_pr = table._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size_eighths)
        b.set(qn('w:color'), hex_color)
        borders.append(b)
    tbl_pr.append(borders)


def _floor_quarter(dt):
    """Rundet ein datetime auf die vorherige volle Viertelstunde ab."""
    return dt.replace(minute=(dt.minute // 15) * 15, second=0, microsecond=0)


def _ceil_quarter(dt):
    """Rundet ein datetime auf die nächste volle Viertelstunde auf."""
    if dt.minute % 15 == 0 and dt.second == 0 and dt.microsecond == 0:
        return dt.replace(microsecond=0)
    floored = dt.replace(minute=(dt.minute // 15) * 15, second=0, microsecond=0)
    return floored + timedelta(minutes=15)


@login_required
def seminar_plan_export(request, course_pk, block_public_id):
    """Erzeugt einen klassischen Wochen-Stundenplan im Word-Format (Schul-Optik):
    pro Kalenderwoche eine Tabelle mit 15-Minuten-Raster links und Mo–Fr als
    Spalten. Vorträge belegen mehrere Slots durch vertikales Cell-Merging.
    """
    from collections import defaultdict
    from datetime import datetime as dt_, timedelta as td_
    from io import BytesIO
    from django.http import HttpResponse
    from django.utils import timezone
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HEADER_FILL      = '1F4E79'
    TIME_COL_FILL    = 'D9E2F3'
    HOUR_LINE_FILL   = 'C5D4EA'
    CELL_CONFIRMED   = 'D5E8D4'
    CELL_PENDING     = 'FFF2CC'
    CELL_DECLINED    = 'F8CECC'
    EMPTY_CELL_FILL  = 'FCFCFC'
    BORDER_COLOR     = '2C3E50'

    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(
        ScheduleBlock, public_id=block_public_id, course=course, block_type='seminar'
    )
    lectures = list(block.lectures.all().order_by('start_datetime'))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    doc.add_heading(f'Stundenplan: {block.name}', level=1)
    intro = doc.add_paragraph()
    intro.add_run('Kurs: ').bold = True
    intro.add_run(course.title)
    intro.add_run('     •     ')
    intro.add_run('Zeitraum: ').bold = True
    intro.add_run(
        f'{block.start_date.strftime("%d.%m.%Y")} – {block.end_date.strftime("%d.%m.%Y")}'
    )
    if block.location:
        intro.add_run('     •     ')
        intro.add_run('Ort: ').bold = True
        intro.add_run(block.location)

    if not lectures:
        doc.add_paragraph('Es sind keine Vorträge angelegt.')
    else:
        by_week = defaultdict(list)
        for lec in lectures:
            local_start = timezone.localtime(lec.start_datetime)
            iso = local_start.isocalendar()
            by_week[(iso.year, iso.week)].append((local_start, lec))

        DUMMY = date(2000, 1, 3)  # ein Montag, irrelevant – nur zum Subtrahieren

        for (iso_year, iso_week) in sorted(by_week.keys()):
            entries = by_week[(iso_year, iso_week)]
            first_local, _ = entries[0]
            monday = (first_local - timedelta(days=first_local.weekday())).date()
            week_dates = [monday + timedelta(days=i) for i in range(5)]

            # ── Zeit-Bounds (gerundet auf Viertelstunden) ────────────────────
            min_dt = None
            max_dt = None
            for local_start, lec in entries:
                local_end = timezone.localtime(lec.end_datetime)
                s_dt = dt_.combine(DUMMY, local_start.time())
                e_dt = dt_.combine(DUMMY, local_end.time())
                if min_dt is None or s_dt < min_dt:
                    min_dt = s_dt
                if max_dt is None or e_dt > max_dt:
                    max_dt = e_dt
            min_dt = _floor_quarter(min_dt)
            max_dt = _ceil_quarter(max_dt)
            total_minutes = int((max_dt - min_dt).total_seconds() // 60)
            num_slots = total_minutes // 15

            # ── Vorträge auf Slot-Indizes mappen ─────────────────────────────
            # cell_map[(weekday, slot_idx)] = (lecture, span_in_slots)
            cell_map = {}
            covered = set()  # (weekday, slot_idx) für gemergte Folge-Zellen
            for local_start, lec in entries:
                local_end = timezone.localtime(lec.end_datetime)
                lec_start_dt = _floor_quarter(dt_.combine(DUMMY, local_start.time()))
                lec_end_dt   = _ceil_quarter(dt_.combine(DUMMY, local_end.time()))
                slot_start = int((lec_start_dt - min_dt).total_seconds() // 60) // 15
                span = int((lec_end_dt - lec_start_dt).total_seconds() // 60) // 15
                wd = local_start.weekday()
                cell_map[(wd, slot_start)] = (lec, max(1, span))
                for k in range(1, max(1, span)):
                    covered.add((wd, slot_start + k))

            # ── Wochen-Überschrift ───────────────────────────────────────────
            doc.add_paragraph()
            wk_heading = doc.add_paragraph()
            run = wk_heading.add_run(
                f'KW {iso_week} – {monday.strftime("%d.%m.")} bis '
                f'{week_dates[4].strftime("%d.%m.%Y")}'
            )
            run.bold = True
            run.font.size = Pt(13)

            # ── Tabelle (Header + num_slots Slots) ───────────────────────────
            table = doc.add_table(rows=1 + num_slots, cols=6)
            table.autofit = False
            _docx_set_table_borders(table, hex_color=BORDER_COLOR, size_eighths='8')

            # Cell-Refs vor Mergen sammeln
            data_cells = []
            for r in range(1, 1 + num_slots):
                data_cells.append([table.rows[r].cells[c] for c in range(6)])

            # Header-Zeile (kompakt, exakte Höhe)
            def _hdr_compact(p):
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0

            hdr_row = table.rows[0]
            hdr_row.height = Cm(0.85)
            hdr_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            hdr = hdr_row.cells
            for c in hdr:
                _docx_set_cell_shading(c, HEADER_FILL)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hdr[0].text = ''
            p_h0 = hdr[0].paragraphs[0]
            p_h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _hdr_compact(p_h0)
            r_h0 = p_h0.add_run('Zeit')
            r_h0.bold = True
            r_h0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r_h0.font.size = Pt(9)

            weekday_names = ['Montag', 'Dienstag', 'Mittwoch', 'Donnerstag', 'Freitag']
            for i, name in enumerate(weekday_names, start=1):
                cell = hdr[i]
                cell.text = ''
                p1 = cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _hdr_compact(p1)
                r1 = p1.add_run(name)
                r1.bold = True
                r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r1.font.size = Pt(9)
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _hdr_compact(p2)
                r2 = p2.add_run(week_dates[i - 1].strftime('%d.%m.%Y'))
                r2.font.size = Pt(7)
                r2.font.color.rgb = RGBColor(0xDD, 0xE6, 0xF2)

            # Spaltenbreiten
            time_w = Cm(1.8)
            day_w = Cm(4.9)
            for row in table.rows:
                row.cells[0].width = time_w
                for i in range(1, 6):
                    row.cells[i].width = day_w

            # ── Zeitspalte + leere Zellen-Hintergrundfarbe ───────────────────
            for slot_idx in range(num_slots):
                slot_dt = min_dt + td_(minutes=15 * slot_idx)
                slot_time = slot_dt.time()
                cell_time = data_cells[slot_idx][0]
                _docx_set_cell_shading(cell_time, TIME_COL_FILL)
                cell_time.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_time.text = ''
                p = cell_time.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                if slot_time.minute == 0:
                    # volle Stunde: fett
                    r = p.add_run(slot_time.strftime('%H:%M'))
                    r.bold = True
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                else:
                    # Viertelstunde: klein und dezent
                    r = p.add_run(slot_time.strftime(':%M'))
                    r.font.size = Pt(6)
                    r.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

                # Reihen-Höhe exakt fixieren – sonst streckt Word die Zeile am Inhalt
                _row = table.rows[1 + slot_idx]
                _row.height = Cm(0.42)
                _row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                # Tagesspalten initial: leerer Hintergrund je nach Slot
                for d in range(5):
                    cell = data_cells[slot_idx][1 + d]
                    cell.text = ''
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if (d, slot_idx) in covered or (d, slot_idx) in cell_map:
                        continue  # wird gleich mit Vortrag gefüllt / gemerged
                    # Volle-Stunden-Linie subtil hervorheben
                    if slot_time.minute == 0:
                        _docx_set_cell_shading(cell, HOUR_LINE_FILL)
                    else:
                        _docx_set_cell_shading(cell, EMPTY_CELL_FILL)

            # ── Vortragsinhalte schreiben ────────────────────────────────────
            def _compact(p):
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0

            for (wd, slot_idx), (lec, span) in cell_map.items():
                cell = data_cells[slot_idx][1 + wd]
                if lec.is_confirmed:
                    _docx_set_cell_shading(cell, CELL_CONFIRMED)
                elif lec.is_declined:
                    _docx_set_cell_shading(cell, CELL_DECLINED)
                else:
                    _docx_set_cell_shading(cell, CELL_PENDING)

                cell.text = ''
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                local_start = timezone.localtime(lec.start_datetime)
                local_end   = timezone.localtime(lec.end_datetime)

                p_time = cell.paragraphs[0]
                p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p_time)
                r_time = p_time.add_run(
                    f'{local_start.strftime("%H:%M")}–{local_end.strftime("%H:%M")}'
                )
                r_time.font.size = Pt(7)
                r_time.font.color.rgb = RGBColor(0x4A, 0x5C, 0x6F)

                p_topic = cell.add_paragraph()
                p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p_topic)
                r_topic = p_topic.add_run(lec.topic)
                r_topic.bold = True
                r_topic.font.size = Pt(8)

                p_speaker = cell.add_paragraph()
                p_speaker.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p_speaker)
                r_speaker = p_speaker.add_run(lec.speaker_name)
                r_speaker.italic = True
                r_speaker.font.size = Pt(7)

                if lec.location:
                    p_loc = cell.add_paragraph()
                    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _compact(p_loc)
                    r_loc = p_loc.add_run(lec.location)
                    r_loc.font.size = Pt(7)
                    r_loc.font.color.rgb = RGBColor(0x4A, 0x5C, 0x6F)

                if lec.is_declined:
                    p_status = cell.add_paragraph()
                    p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _compact(p_status)
                    r_status = p_status.add_run('— abgelehnt —')
                    r_status.font.size = Pt(7)
                    r_status.font.color.rgb = RGBColor(0xB0, 0x00, 0x30)
                elif lec.is_pending:
                    p_status = cell.add_paragraph()
                    p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _compact(p_status)
                    r_status = p_status.add_run('ausstehend')
                    r_status.font.size = Pt(7)
                    r_status.font.color.rgb = RGBColor(0x9C, 0x6F, 0x00)

            # ── Vertikales Mergen für Vorträge mit span > 1 ──────────────────
            for (wd, slot_idx), (lec, span) in cell_map.items():
                if span <= 1:
                    continue
                top = data_cells[slot_idx][1 + wd]
                for k in range(1, span):
                    if slot_idx + k >= num_slots:
                        break
                    top.merge(data_cells[slot_idx + k][1 + wd])

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f'Stand: {date.today().strftime("%d.%m.%Y")}     •     '
            f'Vorträge gesamt: {len(lectures)}     •     '
            f'Bestätigt: {sum(1 for l in lectures if l.is_confirmed)}     •     '
            f'Ausstehend: {sum(1 for l in lectures if l.is_pending)}     •     '
            f'Abgelehnt: {sum(1 for l in lectures if l.is_declined)}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f'Stundenplan_{block.name}_{date.today().strftime("%Y%m%d")}.docx'.replace(' ', '_')
    response = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def lecture_delete(request, course_pk, block_public_id, lecture_public_id):
    course = get_object_or_404(Course, pk=course_pk)
    block = get_object_or_404(
        ScheduleBlock, public_id=block_public_id, course=course, block_type='seminar'
    )
    lecture = get_object_or_404(SeminarLecture, public_id=lecture_public_id, schedule_block=block)
    if request.method == 'POST':
        topic = lecture.topic
        if lecture.status != LECTURE_STATUS_DECLINED:
            from services.notifications import notify_lecture_cancelled
            notify_lecture_cancelled(request, lecture)
        lecture.delete()
        messages.success(request, f'Vortrag „{topic}" wurde gelöscht.')
        return redirect('course:seminar_calendar', course_pk=course.pk,
                        block_public_id=block.public_id)
    return render(request, 'course/lecture_confirm_delete.html', {
        'course': course,
        'schedule_block': block,
        'lecture': lecture,
    })
