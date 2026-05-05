# SPDX-License-Identifier: EUPL-1.2
# SPDX-FileCopyrightText: 2026 devNicoLax
"""
Handbuch-Generator für das Azubi-Portal.
Erzeugt eine vollständige Word-Datei (.docx) mit Installationsanleitung,
Rollenbeschreibungen, Funktionen und Workflows.

Aufruf:
    python generate_handbuch.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ── Farben ────────────────────────────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x1a, 0x56, 0xdb)   # Blau
COLOR_DARK      = RGBColor(0x11, 0x18, 0x27)   # Fast Schwarz
COLOR_GRAY      = RGBColor(0x6b, 0x72, 0x80)   # Grau
COLOR_LIGHT     = RGBColor(0xf3, 0xf4, 0xf6)   # Helles Grau
COLOR_SUCCESS   = RGBColor(0x05, 0x7a, 0x55)   # Grün
COLOR_WARNING   = RGBColor(0xc2, 0x77, 0x03)   # Gelb
COLOR_DANGER    = RGBColor(0x9b, 0x1c, 0x1c)   # Rot
COLOR_WHITE     = RGBColor(0xff, 0xff, 0xff)


def set_cell_bg(cell, color: RGBColor):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table_header_row(table, headers, bg_color=None):
    """Formatiert die erste Zeile einer Tabelle als Kopfzeile."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        if bg_color:
            set_cell_bg(cell, bg_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = COLOR_WHITE if bg_color else COLOR_DARK
                run.font.size = Pt(9)


def style_table(table):
    """Grundformatierung für alle Tabellen."""
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(9)
            if i == 0:
                pass  # Header handled separately
            elif i % 2 == 0:
                set_cell_bg(cell, RGBColor(0xf9, 0xfa, 0xfb))


def add_row(table, *values):
    """Fügt eine Zeile zur Tabelle hinzu."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(9)
    return row


# ── Dokument initialisieren ───────────────────────────────────────────────────

doc = Document()

# Seitenränder
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Standardschriftart
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Überschriften-Stile anpassen
for i, (size, bold, color) in enumerate([
    (22, True, COLOR_PRIMARY),   # Heading 1
    (16, True, COLOR_DARK),      # Heading 2
    (13, True, COLOR_DARK),      # Heading 3
    (11, True, COLOR_GRAY),      # Heading 4
], start=1):
    h = doc.styles[f'Heading {i}']
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = color
    h.font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(16 if i == 1 else 12)
    h.paragraph_format.space_after  = Pt(6 if i <= 2 else 4)


def h1(text): return doc.add_heading(text, level=1)
def h2(text): return doc.add_heading(text, level=2)
def h3(text): return doc.add_heading(text, level=3)
def h4(text): return doc.add_heading(text, level=4)


def para(text, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def info_box(text, color=COLOR_PRIMARY):
    """Einfacher Hinweiskasten als einzeiliger Absatz."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'ℹ  {text}')
    run.font.color.rgb = color
    run.font.italic = True
    run.font.size = Pt(9.5)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_DARK
    p.paragraph_format.space_after = Pt(2)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# TITELSEITE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Azubi-Portal')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Benutzer- und Administrationshandbuch')
run.font.size = Pt(18)
run.font.color.rgb = COLOR_GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Version 1.0  ·  Stand: {datetime.date.today().strftime("%d.%m.%Y")}')
run.font.size = Pt(11)
run.font.color.rgb = COLOR_GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Dieses Handbuch beschreibt Installation, Konfiguration, Rollen, '
    'Funktionen und Workflows des Azubi-Portals – '
    'der zentralen Verwaltungsplattform für Berufsausbildungen.'
)
run.font.size = Pt(11)
run.font.color.rgb = COLOR_GRAY
run.font.italic = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 1: ÜBERBLICK
# ═══════════════════════════════════════════════════════════════════════════════

h1('1  Überblick')

para(
    'Das Azubi-Portal ist eine webbasierte Verwaltungsplattform für die vollständige '
    'Betreuung von Nachwuchskräften in der Berufsausbildung. Es bildet den gesamten '
    'Lebenszyklus einer Ausbildung digital ab – von der Kursplanung über Praktikumseinsätze, '
    'Beurteilungen und Fehlzeitenmanagement bis hin zur abschließenden Archivierung.'
)

h2('1.1  Zielgruppe')
para(
    'Das System richtet sich an Organisationen, die Berufsausbildungen in eigener Regie '
    'durchführen und koordinieren – typischerweise Behörden, öffentliche Einrichtungen '
    'oder größere Unternehmen mit einem strukturierten Ausbildungsprogramm.'
)

h2('1.2  Kernfunktionen im Überblick')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Modul', 'Funktion', 'Nutzergruppe'], COLOR_PRIMARY)
modules = [
    ('Nachwuchskräfte', 'Stammdaten, Status, Kontakthistorie, Notizen, Checklisten', 'Referat, Leitung'),
    ('Kurse & Blöcke', 'Ausbildungspläne, Ablaufblöcke, Kapazitätsplanung', 'Referat, Leitung'),
    ('Praktikumseinsätze', 'Einsatzplanung, Genehmigungsworkflow, Einsatzbriefe', 'Koordination, Referat'),
    ('Ausbildungsnachweise', 'Digitales Berichtsheft, Tageseinträge, Prüfworkflow', 'Azubi, Koordination'),
    ('Beurteilungen', 'Praxistutoren- & Selbstbeurteilungen, Token-Zugang', 'Praxistutor, Koordination'),
    ('Abwesenheiten', 'Urlaub, Krankmeldungen, Abwesenheitsampel', 'Referat, Urlaubsstelle'),
    ('Lerntage', 'Kontingent, Richtlinien, Sperrzeiten, Genehmigung', 'Referat, Azubi'),
    ('Maßnahmen', 'Interventionsprotokoll, Eskalationskette', 'Referat, Leitung'),
    ('Wohnheim', 'Zimmerverwaltung, Belegungsplanung', 'Hausverwaltung'),
    ('Inventar', 'Ausgaben, Rückgaben, QR-Code-Belege', 'Referat'),
    ('Portal', 'Selbstbedienung für Azubis', 'Nachwuchskraft'),
    ('Dokumente', 'Vorlagen, PDF-Generierung, Paperless-Archiv', 'Alle Rollen'),
]
for row in modules:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 2: SYSTEMANFORDERUNGEN & INSTALLATION
# ═══════════════════════════════════════════════════════════════════════════════

h1('2  Systemanforderungen & Installation')

h2('2.1  Voraussetzungen')

para('Für den Betrieb des Azubi-Portals werden folgende Komponenten benötigt:')
bullet('Docker Engine ≥ 24.0')
bullet('Docker Compose ≥ 2.20')
bullet('Mindestens 2 GB RAM, 10 GB freier Festplattenspeicher')
bullet('Netzwerkzugang (für E-Mail-Versand via SMTP)')
bullet('Optional: Paperless-ngx-Instanz für Dokumentenarchivierung')

h2('2.2  Architektur')

para(
    'Das System besteht aus fünf Docker-Containern, die über Docker Compose orchestriert werden:'
)

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Container', 'Image', 'Funktion'], COLOR_DARK)
services = [
    ('db', 'postgres:17-alpine', 'PostgreSQL-Datenbank (persistente Datenhaltung)'),
    ('redis', 'redis:7-alpine', 'Message Broker für Celery-Aufgaben'),
    ('app', 'Azubi (Django/Gunicorn)', 'Webserver, Port 80 → 8000 (3 Worker)'),
    ('celery_worker', 'Azubi (Celery)', 'Asynchrone Hintergrundaufgaben (2 Worker)'),
    ('celery_beat', 'Azubi (Celery Beat)', 'Zeitgesteuerte Aufgaben (Scheduler)'),
]
for row in services:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('2.3  Schnellstart')

para('1.  Repository klonen oder Projektordner entpacken:')
code('git clone <repository-url> azubi-portal')
code('cd azubi-portal')

para('2.  Umgebungsdatei anlegen:')
code('cp .env.example .env')
code('# .env mit einem Editor öffnen und anpassen (siehe Abschnitt 2.4)')

para('3.  Container starten:')
code('docker compose up -d')

para('4.  Datenbank initialisieren (beim ersten Start automatisch):')
code('# Migrations werden beim App-Start automatisch ausgeführt')

para('5.  Superuser anlegen:')
code('docker compose exec app python manage.py createsuperuser')

para('6.  Browser öffnen:')
code('http://<server-ip>/')

info_box(
    'Der erste Start kann 30–60 Sekunden dauern, bis die Datenbankverbindung '
    'aufgebaut ist und Migrationen abgeschlossen sind.'
)

h2('2.4  Konfiguration (.env)')

para(
    'Alle sensiblen Einstellungen werden über eine .env-Datei im Projektverzeichnis '
    'konfiguriert. Diese Datei darf NICHT in die Versionsverwaltung eingecheckt werden.'
)

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Variable', 'Beispielwert', 'Beschreibung'], COLOR_DARK)
env_vars = [
    # Django
    ('SECRET_KEY', '50 zufällige Zeichen', 'Kryptografischer Schlüssel – niemals teilen!'),
    ('DEBUG', 'False', 'Im Produktivbetrieb immer False'),
    ('ALLOWED_HOSTS', 'portal.meinorg.de', 'Kommagetrennte erlaubte Hostnamen'),
    # Datenbank
    ('DB_NAME', 'azubi', 'Name der PostgreSQL-Datenbank'),
    ('DB_USER', 'azubi', 'Datenbankbenutzer'),
    ('DB_PASSWORD', 'sicheres_passwort', 'Datenbankpasswort'),
    ('DB_HOST', 'db', 'Hostname (im Docker-Netz: "db")'),
    ('DB_PORT', '5432', 'PostgreSQL-Port'),
    # E-Mail
    ('EMAIL_HOST', 'mail.meinorg.de', 'SMTP-Server'),
    ('EMAIL_PORT', '587', 'SMTP-Port'),
    ('EMAIL_HOST_USER', 'no-reply@meinorg.de', 'SMTP-Benutzername'),
    ('EMAIL_HOST_PASSWORD', 'passwort', 'SMTP-Passwort'),
    ('EMAIL_USE_TLS', 'True', 'TLS aktivieren (empfohlen)'),
    ('EMAIL_USE_SSL', 'False', 'SSL (nicht zusammen mit TLS)'),
    ('DEFAULT_FROM_EMAIL', 'Azubi-Portal <no-reply@meinorg.de>', 'Absenderadresse'),
    ('DEFAULT_REPLY_TO_EMAIL', 'ausbildung@meinorg.de', 'Antwortadresse'),
    # Redis / Celery
    ('CELERY_BROKER_URL', 'redis://redis:6379/0', 'Redis-Verbindung für Celery'),
    ('CELERY_RESULT_BACKEND', 'redis://redis:6379/0', 'Redis-Verbindung für Ergebnisse'),
    # Paperless
    ('PAPERLESS_URL', 'http://paperless:8000', 'URL der Paperless-ngx-Instanz'),
    ('PAPERLESS_API_KEY', 'api_key_hier', 'API-Schlüssel aus Paperless-Einstellungen'),
    # Site
    ('SITE_BASE_URL', 'https://portal.meinorg.de', 'Basis-URL für E-Mail-Links'),
]
for row in env_vars:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('2.5  Updates')

para('Um das System zu aktualisieren:')
numbered('Neuen Stand herunterladen (git pull oder Archiv entpacken)')
numbered('Container neu bauen und starten:')
code('docker compose build app && docker compose up -d')
numbered('Migrationen werden automatisch beim App-Start ausgeführt')
numbered('Auf Fehlermeldungen in den Logs achten:')
code('docker compose logs -f app')

h2('2.6  Erstkonfiguration nach dem Start')

para(
    'Nach dem ersten Login sollten unter /admin/ folgende Einstellungen vorgenommen '
    'werden, bevor Nachwuchskräfte importiert werden:'
)
numbered('Eigenes Tagesarbeitskonto anlegen und der Gruppe "ausbildungsleitung" zuweisen.')
numbered('SiteConfiguration: Basis-URL, Anonymisierungszeitplan, Erinnerungszeiten setzen.')
numbered('AbsenceSettings: E-Mail der Urlaubsstelle, Bundesland für Feiertage, Word-Vorlage Urlaubsbestätigung.')
numbered('NotificationTemplates: Alle E-Mail-Vorlagen prüfen und an Organisation anpassen.')
numbered('Stammdaten in dieser Reihenfolge: Berufsbilder → Organisationsstruktur → Standorte → Praxistutoren → Koordinationsgruppen → Bewertungsvorlagen → Wohnheime → Inventar.')
numbered('Erst danach: Nachwuchskräfte per CSV/XLSX-Import oder manuell anlegen.')

h2('2.7  HTTPS / Reverse Proxy')

para(
    'Die App lauscht intern nur auf HTTP. Für den produktiven Betrieb wird zwingend '
    'ein Reverse Proxy mit TLS davor­geschaltet (z. B. Caddy, nginx oder Traefik).'
)

h3('Caddy (empfohlen, automatisches Let\'s Encrypt)')
code('azubi.deine-domain.de {\n    encode zstd gzip\n    reverse_proxy localhost:80\n}')

para('Anschließend in der .env die HTTPS-URL eintragen und App neu starten:')
code('CSRF_TRUSTED_ORIGINS=https://azubi.deine-domain.de\nSITE_BASE_URL=https://azubi.deine-domain.de')
code('docker compose restart app')

h2('2.8  Backup')

para('Für ein vollständiges Backup-Konzept (Stufen, GFS-Rotation, Off-Site, Restore-Tests) siehe das gesonderte Backup-Dokument. Mindestens vor Produktivstart:')
bullet('BACKUP_DIR auf eigener Partition / Volume konfigurieren')
bullet('BACKUP_ALERT_EMAILS auf eine regelmäßig gelesene Adresse')
bullet('Ersten manuellen Lauf ausführen: docker compose exec app python manage.py backup_now')
bullet('Empfohlen: restic-Off-Site-Spiegelung auf NAS oder S3-Storage')

h2('2.9  Updates')

numbered('Vorab Backup auslösen: docker compose exec app python manage.py backup_now')
numbered('Neuen Stand holen: git pull')
numbered('Neu bauen und starten: docker compose up -d --build')
numbered('Logs auf Migrationsfehler prüfen: docker compose logs -f app')

h2('2.10  Troubleshooting')

para('Die häufigsten Startprobleme und ihre Lösung:')

t = doc.add_table(rows=1, cols=2)
add_table_header_row(t, ['Symptom', 'Maßnahme'], COLOR_DARK)
ts = [
    ('App startet nicht / Wartet auf DB',
     'docker compose logs db prüfen. Häufig: DB_PASSWORD geändert nach Erststart → entweder altes Passwort wieder setzen oder Volume löschen (Achtung: Daten weg).'),
    ('Migrationsfehler',
     'docker compose exec app python manage.py migrate --plan zeigt geplante Migrationen. Niemals --fake setzen – Fehler an Entwicklerteam weitergeben.'),
    ('E-Mails kommen nicht an',
     'EMAIL_USE_TLS und EMAIL_USE_SSL nicht beide True. Ausgehender Port 587/465 in Firewall offen. Test über Django-Shell mit send_mail().'),
    ('Celery-Aufgaben laufen nicht',
     'Server-Zeitzone prüfen (timedatectl). Nach Korrektur: docker compose restart celery_beat.'),
    ('Paperless-Upload schlägt fehl',
     'PAPERLESS_API_KEY und PAPERLESS_URL prüfen. Token aus Paperless-Einstellungen → API-Token erzeugen.'),
    ('Static Files (CSS/JS) fehlen',
     'docker compose exec app python manage.py collectstatic --noinput; docker compose restart app'),
]
for row in ts:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('2.11  Härtungs-Checkliste für Produktion')

bullet('DEBUG=False in .env (zwingend)')
bullet('HTTPS-only — Reverse Proxy mit aktuellem TLS-Zertifikat')
bullet('ALLOWED_HOSTS exakt eingeschränkt — kein Sternchen')
bullet('SECRET_KEY mind. 50 Zeichen, einzigartig, niemals geloggt')
bullet('DB-Passwort mind. 24 Zeichen, kein Wiederverwendungs-Passwort')
bullet('SMTP-Zugangsdaten an dedizierten Account binden, kein persönliches Postfach')
bullet('Server-Firewall: nur 80/443 öffentlich, SSH auf Admin-IPs eingeschränkt')
bullet('Backups laufen, Alert-Mail wird regelmäßig gelesen')
bullet('Quartalsweiser Restore-Test im Audit-Log sichtbar')
bullet('Off-Site-Spiegelung (restic) eingerichtet, restic.key an zwei sicheren Orten')
bullet('Server-Updates automatisiert (z. B. unattended-upgrades)')
bullet('Monitoring auf /healthz (Liveness) und /readyz (Readiness inkl. DB/Cache) – Uptime Robot, Prometheus Blackbox o. Ä.')
bullet('Logging zentralisiert (z. B. Loki/Promtail oder Filebeat)')


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 3: ROLLEN & BERECHTIGUNGEN
# ═══════════════════════════════════════════════════════════════════════════════

h1('3  Rollen & Berechtigungen')

para(
    'Das Azubi-Portal verwendet ein rollenbasiertes Zugriffssystem. Jeder Benutzer '
    'hat genau eine Primärrolle, die seinen Zugang zu Menüpunkten und Funktionen '
    'steuert. Die Zuweisung erfolgt über das Django-Adminpanel oder die '
    'Systemkonfiguration.'
)

h2('3.1  Rollenübersicht')

t = doc.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Rolle', 'Technischer Name', 'Primäre Aufgabe', 'Zugang'], COLOR_PRIMARY)
roles = [
    ('Ausbildungsleitung', 'is_training_director', 'Strategische Steuerung, Berichte, Konfiguration', 'Vollzugriff'),
    ('Ausbildungsreferat', 'is_training_office', 'Operative Verwaltung aller Azubis und Prozesse', 'Fast vollständig'),
    ('Ausbildungskoordination', 'is_training_coordinator', 'Einsatzplanung, Praxistutoren, Beurteilungen', 'Eingeschränkt'),
    ('Ausbildungsverantwortliche', 'is_training_responsible', 'Lesezugriff auf Azubidaten und Nachweise', 'Lesend'),
    ('Hausverwaltung', 'is_dormitory_management', 'Zimmerbelegung und Wohnheimplanung', 'Wohnheim'),
    ('Reisekostenstelle', 'is_travel_expense_office', 'Einsicht in Azubi- und Kursdaten', 'Eingeschränkt lesend'),
    ('Nachwuchskraft (Azubi)', '(Portal-User)', 'Selbstbedienung über das Portal', 'Nur eigene Daten'),
]
for row in roles:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('3.2  Detaillierte Berechtigungen')

h3('Ausbildungsleitung')
para(
    'Die Ausbildungsleitung hat nahezu vollständigen Zugriff auf alle Bereiche des '
    'Systems. Darüber hinaus stehen exklusiv zur Verfügung:'
)
bullet('Systemweite Konfiguration (Marke, Farben, Erinnerungszeiten, Datenschutztexte)')
bullet('Verwaltung aller Benachrichtigungsvorlagen')
bullet('Berichte und Auswertungen über alle Azubis und Kurse')
bullet('Verwaltung von Abwesenheitseinstellungen (Bundesland, Urlaubsstellen-E-Mail)')
bullet('Maßnahmen-/Interventionsverwaltung')
bullet('Datenschutz: manuelle und automatische Anonymisierung')
bullet('Dokumentenvorlagen-Verwaltung für alle Module')

h3('Ausbildungsreferat')
para('Das Ausbildungsreferat führt die operative Tagesarbeit durch:')
bullet('Alle Nachwuchskraft-Stammdaten bearbeiten (inkl. benutzerdefinierter Felder)')
bullet('Kurse und Ablaufblöcke anlegen und verwalten')
bullet('Urlaubsanträge genehmigen oder ablehnen')
bullet('Krankmeldungen erfassen und schließen')
bullet('Lerntage genehmigen oder ablehnen')
bullet('Checklisten und Onboarding-Aufgaben verwalten')
bullet('Inventar ausgeben und zurücknehmen')
bullet('Wohnheimbelegungen zuweisen')
bullet('Signierte PDFs für Urlaubsgenehmigungen und Beurteilungen erstellen')

h3('Ausbildungskoordination')
para('Die Koordination ist für Praktikumseinsätze zuständig:')
bullet('Übersicht aller betreuten Nachwuchskräfte (lesen)')
bullet('Praktikumseinsätze für zugeordnete Organisationseinheiten anlegen')
bullet('Einsatzbriefe generieren und versenden')
bullet('Praxistutoren verwalten (anlegen, bestätigen)')
bullet('Beurteilungen anfordern, prüfen und bestätigen')
bullet('Einsätze genehmigen oder ablehnen (als Chief Instructor)')
bullet('Eigene Kursübersicht und Kalender einsehen')

h3('Nachwuchskraft (Portal)')
para('Die Nachwuchskraft hat ausschließlich Zugriff auf eigene Daten:')
bullet('Dashboard mit aktuellem Einsatz und anstehenden Terminen')
bullet('Eigene Stammdaten einsehen')
bullet('Stationsplan (Übersicht aller Praktikumseinsätze)')
bullet('Ausbildungsnachweise / Berichtsheft führen')
bullet('Urlaubsanträge stellen und stornieren')
bullet('Lerntage beantragen')
bullet('Selbstbeurteilungen ausfüllen')
bullet('Beurteilungen durch Praxistutoren einsehen')
bullet('Persönlichen Kalender einsehen')

h2('3.3  Benutzerverwaltung')

para('Neue Benutzerkonten werden im Django-Adminpanel angelegt (/admin/):')
numbered('Unter „Benutzer" → „Benutzer hinzufügen" ein Konto erstellen')
numbered('Passwort setzen und Benutzerprofil befüllen (Name, Abteilung, Telefon)')
numbered('Unter „Services" → „Ausbildungsreferat-Profile" die Rolle zuweisen')
numbered('Für Portal-Zugang: Unter dem Azubi-Datensatz den Portal-User verknüpfen')

info_box(
    'Rollenprofile werden im Admin unter „Services → Ausbildungsreferat Profile" '
    'verwaltet. Hier werden auch Berechtigungen für Inventar, Abwesenheit und '
    'Lerntage pro Nutzer granular gesteuert.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 4: NACHWUCHSKRÄFTEVERWALTUNG
# ═══════════════════════════════════════════════════════════════════════════════

h1('4  Nachwuchskräfteverwaltung')

h2('4.1  Stammdaten')

para(
    'Jede Nachwuchskraft wird mit einem vollständigen Datensatz erfasst. '
    'Die Eingabe erfolgt über das Formular unter Nachwuchskräfte → Neu anlegen.'
)

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Feld', 'Beschreibung'], COLOR_DARK)
fields = [
    ('Vorname / Nachname', 'Vollständiger Name'),
    ('Geburtsdatum / Geburtsort', 'Pflichtangabe für offizielle Dokumente'),
    ('Anschrift', 'Straße, Hausnummer, PLZ, Ort'),
    ('Telefon', 'Private Telefonnummer'),
    ('E-Mail privat', 'Private E-Mail-Adresse'),
    ('E-Mail Kennung', 'Dienstliche E-Mail / Portalzugang'),
    ('Geschlecht / Anrede', 'Für korrekte Ansprache in E-Mails und Dokumenten'),
    ('Kurs', 'Zuweisung zu einem Ausbildungskurs'),
    ('Beschäftigungsverhältnis', 'Art der Beschäftigung'),
    ('Status', 'Aktueller Ausbildungsstatus (konfigurierbar, mit Farbe)'),
    ('Benutzerdefinierte Felder', 'Frei konfigurierbare Zusatzfelder (Text, Zahl, Datum, Ja/Nein)'),
]
for row in fields:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('4.2  Statusverwaltung')

para(
    'Jede Nachwuchskraft trägt einen konfigurierbaren Status. '
    'Statuswechsel werden mit Zeitstempel protokolliert. '
    'Die Farben dienen der schnellen Orientierung in Übersichtslisten:'
)
bullet('Grün (success) – z. B. aktiv in Ausbildung')
bullet('Grau (secondary) – z. B. abgeschlossen, inaktiv')
bullet('Rot (danger) – z. B. abgebrochen, kritisch')
bullet('Gelb (warning) – z. B. pausiert, unklar')
bullet('Blau (info) – z. B. Sonderstatus')

info_box(
    'Status und ihre Bedeutungen können im Adminpanel unter '
    '„Student → Status" frei definiert werden.'
)

h2('4.3  Kontakthistorie')

para(
    'Unter dem Reiter „Kontakthistorie" können Gespräche, Telefonate und '
    'E-Mail-Kontakte mit der Nachwuchskraft protokolliert werden:'
)
bullet('Kontaktart: Telefon, Vor-Ort-Gespräch, E-Mail')
bullet('Datum und Uhrzeit des Kontakts')
bullet('Freitext: Inhalt / Verlauf des Gesprächs')
bullet('Anlass und Reaktion der Nachwuchskraft')

h2('4.4  Interne Notizen')

para(
    'Interne Notizen sind nur für Ausbildungsreferat und Leitung sichtbar. '
    'Sie können angeheftet werden, um wichtige Informationen oben anzuzeigen. '
    'Die Nachwuchskraft hat keinen Zugriff auf diese Notizen.'
)

h2('4.5  Checklisten')

para(
    'Checklisten unterstützen das strukturierte Onboarding. '
    'Aus hinterlegten Vorlagen werden Aufgabenlisten je Nachwuchskraft erstellt:'
)
bullet('Vorlage definiert wiederkehrende Aufgaben (z. B. „Arbeitsausweis beantragt", „IT-Konto eingerichtet")')
bullet('Pro Nachwuchskraft wird eine Instanz erstellt')
bullet('Einzelne Aufgaben werden abhakt und mit Datum versehen')
bullet('Fortschrittsanzeige zeigt erledigte vs. offene Aufgaben')

h2('4.6  Notenübersicht')

para(
    'Unter „Noten" werden Prüfungs- und Leistungsbewertungen erfasst. '
    'Je nach Konfiguration des Berufsbilds sind verschiedene Notentypen '
    '(z. B. Zwischenprüfung, Abschlussprüfung, Modultest) hinterlegt. '
    'Anhänge (z. B. Zeugnisse) können direkt in Paperless-ngx abgelegt werden.'
)

h2('4.7  Benutzerdefinierte Felder')

para(
    'Das System erlaubt es, beliebige Zusatzfelder zu definieren. '
    'Diese gelten für alle Nachwuchskräfte und unterstützen folgende Datentypen:'
)
bullet('Text (einzeilig oder mehrzeilig)')
bullet('Zahl')
bullet('Datum')
bullet('Ja / Nein (Checkbox)')

h2('4.8  Datenschutz & Anonymisierung')

para(
    'Zum Schutz personenbezogener Daten bietet das System zwei Mechanismen:'
)
bullet(
    'Manuelle Anonymisierung: Einzelne Datensätze können durch berechtigte Nutzer '
    'manuell anonymisiert werden (alle personenbezogenen Felder werden überschrieben).'
)
bullet(
    'Automatische Anonymisierung: Nachwuchskräfte, deren Status seit der konfigurierten '
    'Frist (Standard: 12 Monate) nicht mehr geändert wurde, werden automatisch '
    'anonymisiert. Der Zeitpunkt ist täglich per Celery-Beat konfigurierbar.'
)

info_box(
    'Die Frist für die automatische Anonymisierung wird unter '
    'Systemkonfiguration → Anonymisierung eingestellt.'
)

h2('4.9  Import & Export')

para('Massenoperationen werden über CSV/Excel unterstützt:')
bullet('Import: CSV (Semikolon, UTF-8) oder .xlsx – Vorlagen-Pflichtfelder beachten')
bullet('Export: Alle Nachwuchskräfte oder gefilterte Auswahl als CSV')
bullet('Importfehler werden pro Zeile angezeigt, korrekte Zeilen werden dennoch importiert')


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 5: KURSE & AUSBILDUNGSPLÄNE
# ═══════════════════════════════════════════════════════════════════════════════

h1('5  Kurse & Ausbildungspläne')

h2('5.1  Berufsbilder')

para(
    'Berufsbilder (z. B. Verwaltungsfachangestellte/r, IT-Systemkaufmann/-frau) '
    'bilden die Grundlage für alle ausbildungsspezifischen Konfigurationen:'
)
bullet('Name, Abschlussgrad und gesetzliche Grundlage')
bullet('Laufbahn und Fachrichtung')
bullet('Notentypen (z. B. Zwischenprüfung, Abschlussprüfung)')
bullet('Beurteilungsvorlagen')
bullet('Lerntage-Richtlinien')
bullet('Pflicht zur Führung von Ausbildungsnachweisen (Ja/Nein)')

h2('5.2  Kurse anlegen')

para(
    'Ein Kurs bündelt eine Gruppe von Nachwuchskräften des gleichen Berufsbilds '
    'und Jahrgangs. Pflichtangaben beim Anlegen:'
)
bullet('Bezeichnung des Kurses')
bullet('Zugehöriges Berufsbild')
bullet('Startdatum und Enddatum der Ausbildung')

h2('5.3  Ablaufblöcke (Stationsplan)')

para(
    'Ein Kurs wird in Ablaufblöcke unterteilt. Jeder Block entspricht einer '
    'Phase der Ausbildung (z. B. Einführungslehrgang, Praktikum OE 1, Theoriephasе):'
)
bullet('Name des Blocks')
bullet('Start- und Enddatum')
bullet('Art: Theorietag / Praktikum / Sonstiges')
bullet('Einsatzort (Standort aus dem Organisationsverzeichnis)')
bullet('Ist-Praktikum-Flag: Aktiviert Einsatz- und Beurteilungsworkflows für diesen Block')

h2('5.4  Kurskalender')

para(
    'Der Kurskalender zeigt alle Ablaufblöcke eines Kurses chronologisch. '
    'Von hier aus können Praktikumseinsätze direkt angelegt werden.'
)

h2('5.5  Kapazitätsplanung')

para(
    'Die Kapazitätsansicht zeigt für jeden Block, wie viele Nachwuchskräfte '
    'welchen Organisationseinheiten zugeordnet sind – '
    'und ob die jeweilige Kapazitätsgrenze erreicht ist.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 6: PRAKTIKUMSEINSÄTZE
# ═══════════════════════════════════════════════════════════════════════════════

h1('6  Praktikumseinsätze')

h2('6.1  Workflow im Überblick')

para('Der Workflow eines Praktikumseinsatzes läuft in folgenden Schritten ab:')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Schritt', 'Aktion', 'Verantwortlich'], COLOR_PRIMARY)
steps = [
    ('1', 'Einsatz anlegen (Nachwuchskraft, OE, Zeitraum, Praxistutor)', 'Koordination / Referat'),
    ('2', 'Automatische Benachrichtigung an Chief Instructor', 'System'),
    ('3', 'Chief Instructor genehmigt oder lehnt ab', 'Chief Instructor / Koordination'),
    ('4', 'Einsatzbrief generieren und versenden (optional)', 'Referat'),
    ('5', '7 Tage vor Beginn: Erinnerungs-E-Mail an Praxistutor', 'System (Celery)'),
    ('6', 'Praxistutor erhält Token-Link für Beurteilung', 'System (Celery)'),
    ('7', 'Praxistutor füllt Beurteilung aus', 'Praxistutor'),
    ('8', 'Koordination prüft und bestätigt Beurteilung', 'Koordination'),
    ('9', 'Beurteilung gesperrt, signiertes PDF abrufbar', 'System'),
]
for row in steps:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('6.2  Status eines Einsatzes')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Status', 'Anzeige', 'Bedeutung'], COLOR_DARK)
statuses = [
    ('Ausstehend', 'Gelb', 'Einsatz wurde angelegt, Genehmigung steht aus'),
    ('Angenommen', 'Grün', 'Einsatz wurde durch Chief Instructor genehmigt'),
    ('Abgelehnt', 'Rot', 'Einsatz wurde abgelehnt (Grund wird angezeigt)'),
]
for row in statuses:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('6.3  Einsatzbriefe')

para(
    'Für genehmigte Einsätze können automatisch Zuweisungsschreiben (.docx) '
    'auf Basis hinterlegter Vorlagen generiert werden. '
    'Die Vorlagen unterstützen Platzhalter wie:'
)
bullet('{{ student_vorname }}, {{ student_nachname }}')
bullet('{{ block_beginn }}, {{ block_ende }}')
bullet('{{ einheit_name }}, {{ standort }}')
bullet('{{ praxistutor_name }}, {{ praxistutor_email }}')

h2('6.4  Einsatz teilen (Split)')

para(
    'Muss ein Einsatz während der Laufzeit auf eine andere Organisationseinheit '
    'aufgeteilt werden, kann über die Funktion „Einsatz teilen" ein Teilungsantrag '
    'gestellt werden. Das System erstellt daraus zwei separate Einsätze.'
)

h2('6.5  Stationsbriefe')

para(
    'Neben Zuweisungsbriefen können Stationsbriefe generiert werden, '
    'die eine Übersicht aller Stationen einer Nachwuchskraft für einen Block bündeln.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 7: AUSBILDUNGSNACHWEISE (BERICHTSHEFT)
# ═══════════════════════════════════════════════════════════════════════════════

h1('7  Ausbildungsnachweise (Digitales Berichtsheft)')

para(
    'Das Modul „Ausbildungsnachweise" bildet das gesetzlich vorgeschriebene '
    'Berichtsheft digital ab. Ob eine Nachwuchskraft Nachweise führen muss, '
    'wird am Berufsbild konfiguriert.'
)

h2('7.1  Wochennachweis erstellen (Azubi)')

para('Die Nachwuchskraft führt wöchentlich einen Nachweis:')
numbered('Im Portal unter „Ausbildungsnachweise" → „Neue Woche"')
numbered('Für jeden der 5 Wochentage eine Tagesaktivität eintragen')
numbered('Aktivitätstypen: Praktische Ausbildung, Berufsschule, Urlaub/Krank/Feiertag')
numbered('Nachweis als Entwurf speichern oder direkt einreichen')

h2('7.2  Status eines Nachweises')

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Status', 'Bedeutung'], COLOR_DARK)
pof_statuses = [
    ('Entwurf', 'Vom Azubi gespeichert, noch nicht eingereicht'),
    ('Eingereicht', 'Zur Prüfung vorgelegt'),
    ('Genehmigt', 'Durch Koordination/Referat freigegeben'),
    ('Abgelehnt', 'Mit Korrekturhinweis zurückgegeben'),
]
for row in pof_statuses:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('7.3  Prüfworkflow (Koordination / Referat)')

para('Eingereichte Nachweise werden wie folgt bearbeitet:')
numbered('Übersicht aller eingereichten Nachweise aufrufen')
numbered('Einzelnachweis öffnen und Tageseinträge prüfen')
numbered('Genehmigen: Nachweis wird als „Genehmigt" markiert')
numbered('Ablehnen: Korrekturhinweis eingeben, Nachweis geht zurück an Azubi')

h2('7.4  Erinnerungen')

para(
    'Täglich um 07:00 Uhr prüft das System, ob Nachwuchskräfte '
    'überfällige oder nicht eingereichte Nachweise haben. '
    'Betroffene Azubis erhalten automatisch eine Erinnerungs-E-Mail.'
)

h2('7.5  PDF-Export')

para(
    'Alle Nachweise einer Nachwuchskraft können als vollständige .docx-Datei '
    'exportiert werden. Dafür muss eine Export-Vorlage hinterlegt sein.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 8: BEURTEILUNGSSYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

h1('8  Beurteilungssystem')

h2('8.1  Beurteilungsvorlagen')

para(
    'Für jedes Berufsbild wird eine Beurteilungsvorlage konfiguriert. '
    'Diese legt fest, welche Kriterien bewertet werden und welche Skala gilt:'
)
bullet('Notenskala 1,0 – 6,0 (Schulnoten)')
bullet('Punkteskala 0 – 15 (Oberstufenpunkte)')
para('Kriterien können in Kategorien gruppiert werden und Hilfetexte enthalten.')

h2('8.2  Beurteilung durch Praxistutoren')

para(
    'Praxistutoren werden per E-Mail mit einem personalisierten Token-Link '
    'zur Beurteilung eingeladen. Für den Zugang ist kein Benutzerkonto erforderlich.'
)

para('Ablauf:')
numbered('System generiert Token und versendet Link an Praxistutor')
numbered('Praxistutor öffnet Link, füllt alle Kriterien aus')
numbered('Optional: Gesamtkommentar erfassen')
numbered('Beurteilung einreichen → Status wechselt auf „Eingereicht"')
numbered('Koordination/Referat prüft und bestätigt')
numbered('Status wechselt auf „Bestätigt" – Beurteilung ist gesperrt')

h2('8.3  Selbstbeurteilung (Azubi)')

para(
    'Parallel zur Praxistutor-Beurteilung kann die Nachwuchskraft '
    'eine Selbstbeurteilung ausfüllen. Diese wird der Koordination '
    'im Vergleich zur Fremdbewertung angezeigt.'
)

h2('8.4  Signiertes PDF')

para(
    'Nach Bestätigung einer Beurteilung kann ein signiertes PDF heruntergeladen werden. '
    'Das Dokument enthält alle Bewertungen sowie einen elektronischen Signaturblock '
    'mit Name, Funktion und Zeitstempel des Bestätigenden '
    '(einfache elektronische Signatur gem. eIDAS Art. 3 Nr. 10).'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 9: ABWESENHEITSVERWALTUNG
# ═══════════════════════════════════════════════════════════════════════════════

h1('9  Abwesenheitsverwaltung')

h2('9.1  Urlaubsanträge')

h3('Antrag stellen (Azubi im Portal)')
numbered('Im Portal unter „Urlaub" → „Neuen Antrag stellen"')
numbered('Zeitraum (Von – Bis) und optionale Anmerkungen eingeben')
numbered('Antrag absenden → Status: Ausstehend')
numbered('E-Mail-Benachrichtigung an das Ausbildungsreferat')

h3('Genehmigungsworkflow (Referat)')
numbered('Urlaubsantrag in der Liste öffnen')
numbered('„Verbindlich genehmigen" oder „Ablehnen" (mit Begründung)')
numbered('Automatische Benachrichtigung an Nachwuchskraft')
numbered('Genehmigte Anträge: Signiertes PDF abrufbar')
numbered('Täglicher Batch-Versand genehmigter Anträge an Urlaubsstelle (08:00 Uhr)')

h3('Bearbeitung durch Urlaubsstelle')
para(
    'Die Urlaubsstelle erhält täglich eine E-Mail mit einem einmaligen Token-Link. '
    'Über diesen Link trägt sie den Resturlaub ein und schließt die Anträge ab – '
    'ohne eigenes Benutzerkonto im System.'
)

h3('Status eines Urlaubsantrags')
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Status', 'Bedeutung'], COLOR_DARK)
vstatus = [
    ('Ausstehend', 'Eingereicht, Entscheidung steht aus'),
    ('Genehmigt', 'Durch Referat freigegeben'),
    ('Abgelehnt', 'Abgelehnt, Begründung vorhanden'),
    ('Durch Urlaubsstelle bearbeitet', 'Resturlaub eingetragen, abgeschlossen'),
    ('Storniert', 'Antrag wurde zurückgezogen'),
]
for row in vstatus:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('9.2  Stornierungsanträge')

para(
    'Bereits genehmigte Urlaubsanträge können über einen Stornierungsantrag '
    'zurückgezogen werden. Stornierungen durchlaufen denselben Genehmigungsprozess '
    'wie normale Urlaubsanträge.'
)

h2('9.3  Krankmeldungen')

para('Krankmeldungen werden durch das Ausbildungsreferat erfasst:')
bullet('Erfassung von Start und (bei Rückkehr) Ende der Erkrankung')
bullet('Krankheitsart: Selbstauskunft, K-Taste, Attest')
bullet('Automatische Meldung an die Urlaubsstelle (täglich 08:05 Uhr)')

h2('9.4  Abwesenheitsampel')

para(
    'Das System berechnet kontinuierlich die Fehlzeitenquote jeder Nachwuchskraft '
    'in Relation zu den Gesamtarbeitstagen des Kurses:'
)
bullet('Grün: Fehlzeitenquote unter 5 %')
bullet('Gelb: Fehlzeitenquote 5 – 9,9 %')
bullet('Rot: Fehlzeitenquote 10 % und mehr')

para(
    'Bei einem Statuswechsel der Ampel wird automatisch eine Benachrichtigung '
    'an die Ausbildungsleitung versendet. Die Ampel ist im Azubi-Profil sichtbar.'
)

info_box(
    'Das Bundesland für die Feiertagsberechnung wird unter '
    'Abwesenheiten → Einstellungen → Bundesland konfiguriert.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 10: LERNTAGE
# ═══════════════════════════════════════════════════════════════════════════════

h1('10  Lerntage')

para(
    'Das Lerntage-Modul verwaltet Ansprüche und Anträge für Lern- und '
    'Prüfungsvorbereitungstage. Die Richtlinien sind pro Berufsbild konfigurierbar '
    'und ermöglichen eine präzise Steuerung der Kontingente.'
)

h2('10.1  Richtlinien-Konfiguration')

para('Je Berufsbild und Ausbildungsjahr können folgende Parameter gesetzt werden:')
bullet('Geltungsbereich: Nur Praktikumsphasen oder gesamte Ausbildung')
bullet('Kontingentart: Gesamtbudget, pro Block, pro Woche oder pro Monat')
bullet('Jahrweises Kontingent (Jahr 1, Jahr 2, Jahr 3+)')
bullet('Monatliches Maximum (Kappungsgrenze)')
bullet('Übertrag auf Folgemonat: Ja / Nein')
bullet('Prüfungsvorbereitungstage (separates Kontingent)')
bullet('Mindestvorlauf in Tagen für Antragstellung')
bullet('Maximale aufeinanderfolgende Tage pro Antrag')
bullet('Erlaubte Wochentage (z. B. nur Mo–Do)')

h2('10.2  Sperrzeiten')

para(
    'Sperrzeiten schließen bestimmte Zeiträume aus der Antragstellung aus '
    '(z. B. Prüfungswochen, Betriebsferien). Sie werden pro Berufsbild hinterlegt.'
)

h2('10.3  Antragsprozess')

numbered('Nachwuchskraft stellt Antrag im Portal (Datum, Art: Lerntag / Prüfungsvorbereitung)')
numbered('System prüft automatisch: Kontingent, Vorlauf, Sperrzeiten, Wochentag')
numbered('Bei Regelverletzung: sofortige Ablehnung mit Begründung')
numbered('Sonst: Antrag landet beim Ausbildungsreferat zur Entscheidung')
numbered('Referat genehmigt oder lehnt ab → E-Mail-Benachrichtigung an Azubi')


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 11: MAßNAHMEN & INTERVENTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

h1('11  Maßnahmen & Interventionen')

para(
    'Das Interventionsmodul dokumentiert Gespräche und Maßnahmen, '
    'die aufgrund von Auffälligkeiten (Fehlzeiten, Bewertungen, Verhalten) '
    'eingeleitet werden.'
)

h2('11.1  Kategorien')

para(
    'Interventionskategorien werden im Adminpanel definiert. '
    'Jede Kategorie hat eine Eskalationsstufe (1–4) und eine Farbe '
    'zur visuellen Unterscheidung.'
)

h2('11.2  Intervention anlegen')

para('Pflichtfelder beim Anlegen einer Intervention:')
bullet('Auslöser: Fehlzeit, Beurteilung, Verhalten/Disziplin, Sonstiges')
bullet('Verknüpfung mit konkreter Krankmeldung oder Beurteilung (optional)')
bullet('Datum des Gesprächs')
bullet('Anwesende Personen')
bullet('Beschreibung des Sachverhalts')
bullet('Vereinbarte Maßnahmen und Folgetermin')

h2('11.3  Eskalationskette')

para(
    'Eine Intervention kann mit einer nachfolgenden Intervention verknüpft werden. '
    'So entsteht eine lückenlose Eskalationskette. '
    'Das System ermöglicht es, offene Maßnahmen zu verfolgen und abzuschließen.'
)

h2('11.4  Status')
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Status', 'Bedeutung'], COLOR_DARK)
int_statuses = [
    ('Offen', 'Maßnahme eingeleitet, Folgetermin steht aus'),
    ('In Bearbeitung', 'Maßnahmen laufen'),
    ('Geschlossen', 'Erfolgreich abgeschlossen'),
    ('Eskaliert', 'An höhere Stelle übergeben'),
]
for row in int_statuses:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 12: WOHNHEIMVERWALTUNG
# ═══════════════════════════════════════════════════════════════════════════════

h1('12  Wohnheimverwaltung')

para(
    'Das Wohnheimmodul verwaltet die Unterbringung von Nachwuchskräften '
    'in Gemeinschaftsunterkünften.'
)

h2('12.1  Wohnheime & Zimmer')
bullet('Wohnheime mit Adresse und Kontaktdaten anlegen')
bullet('Zimmer mit Kapazität (max. Belegung) zuordnen')
bullet('Zimmer für Renovierungen o. ä. sperren (mit Begründung und Zeitraum)')

h2('12.2  Zimmerbelegung')
numbered('Nachwuchskraft auswählen')
numbered('Zimmer und Zeitraum (Von – Bis) zuweisen')
numbered('Das System prüft Kapazität und Überschneidungen automatisch')
numbered('Reservierungsbestätigung als .docx generierbar')

h2('12.3  Belegungskalender')
para(
    'Der Belegungskalender zeigt tagesgenau, welche Zimmer wann belegt sind, '
    'und unterstützt die Vorausplanung von Unterbringungen.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 13: INVENTARVERWALTUNG
# ═══════════════════════════════════════════════════════════════════════════════

h1('13  Inventarverwaltung')

para(
    'Das Inventarmodul verwaltet die Ausgabe und Rückgabe von Gegenständen '
    '(Dienstausweis, Laptop, Schlüssel usw.) an Nachwuchskräfte.'
)

h2('13.1  Artikel & Kategorien')
bullet('Artikel mit Seriennummer, Status und Kategorie anlegen')
bullet('Kategorien mit individuellen Symbolen und Quittungsvorlagen versehen')
bullet('Statuswerte: Verfügbar, Ausgegeben, Defekt, Ausgemustert')

h2('13.2  Ausgabe & Rückgabe')
numbered('Artikel aus dem Katalog auswählen')
numbered('Nachwuchskraft und Ausgabedatum festlegen')
numbered('Quittung als .docx mit QR-Code generieren und aushändigen')
numbered('Nachwuchskraft unterschreibt Quittung (physisch)')
numbered('Unterschriebene Quittung einscannen und per Mehrfach-Scan hochladen')
numbered('QR-Code auf dem Scan wird automatisch erkannt und der Ausgabe zugeordnet')
numbered('Bei Rückgabe: Rückgabedatum eintragen')

h2('13.3  Mehrfach-Scan (QR-Code-Erkennung)')
para(
    'Mehrere eingescannte Quittungen (PDF) können gleichzeitig hochgeladen werden. '
    'Das System erkennt den QR-Code automatisch, verknüpft das Dokument '
    'mit der richtigen Ausgabe und speichert es in Paperless-ngx.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 14: AZUBI-PORTAL
# ═══════════════════════════════════════════════════════════════════════════════

h1('14  Azubi-Portal (Nachwuchskräfte-Selbstbedienung)')

para(
    'Das Portal ermöglicht Nachwuchskräften, auf ihre eigenen Daten zuzugreifen '
    'und Anträge zu stellen – ohne Eingriff des Referats.'
)

h2('14.1  Dashboard')
bullet('Aktueller und nächster Praktikumseinsatz auf einen Blick')
bullet('Offene Ausbildungsnachweise (falls Pflicht im Berufsbild)')
bullet('Aktuelle Noten')
bullet('Ausstehende Selbstbeurteilungen')
bullet('Vorschau des Ausbildungskalenders')

h2('14.2  Stationsplan')
para(
    'Vollständige chronologische Übersicht aller Praktikumseinsätze '
    'mit Organisationseinheit, Praxistutor und Zeitraum.'
)

h2('14.3  Urlaub')
bullet('Urlaubsanträge stellen und Bearbeitungsstatus verfolgen')
bullet('Genehmigte Anträge einsehen')
bullet('Stornierungsanträge für genehmigte Urlaubsanträge stellen')

h2('14.4  Lerntage')
bullet('Offenes Kontingent einsehen (pro Monat und gesamt)')
bullet('Anträge für Lerntage und Prüfungsvorbereitungstage stellen')
bullet('Status eigener Anträge verfolgen')

h2('14.5  Ausbildungsnachweise')
bullet('Wöchentliche Nachweise anlegen und einreichen')
bullet('Status (Entwurf, Eingereicht, Genehmigt, Abgelehnt) verfolgen')
bullet('Korrekturhinweise bei Ablehnung einsehen')

h2('14.6  Selbstbeurteilung')
bullet('Selbstbeurteilung für abgeschlossene Einsätze ausfüllen')
bullet('Eigene Beurteilungen durch Praxistutoren einsehen (nach Bestätigung)')

h2('14.7  Kalender')
para(
    'Persönlicher Ausbildungskalender mit Ablaufblöcken, '
    'Praktikumseinsätzen und Schultagen.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 15: BENACHRICHTIGUNGEN
# ═══════════════════════════════════════════════════════════════════════════════

h1('15  Benachrichtigungen')

h2('15.1  E-Mail-Benachrichtigungen')

para(
    'Das System versendet automatisch E-Mails bei relevanten Ereignissen. '
    'Alle Vorlagen sind konfigurierbar und unterstützen Platzhalter.'
)

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Ereignis', 'Empfänger', 'Beschreibung'], COLOR_PRIMARY)
notifications = [
    ('Neue Nachwuchskraft zugewiesen', 'Praxistutor', 'Neue Person in Betreuung'),
    ('Praxistutor bestätigt', 'Praxistutor', 'Kontobestätigung'),
    ('Einsatz angelegt/geändert', 'Chief Instructor', 'Neuer oder geänderter Einsatz'),
    ('Chief-Willkommen', 'Chief Instructor', 'Kontozugang und Einweisung'),
    ('Erinnerung Beginn (7 Tage)', 'Praxistutor', 'Bevorstehender Einsatzbeginn'),
    ('Erinnerung Ende (7 Tage)', 'Praxistutor', 'Bevorstehendes Einsatzende'),
    ('Einsatz genehmigt', 'Koordination', 'Genehmigungsentscheidung'),
    ('Einsatz abgelehnt', 'Koordination', 'Ablehnung mit Begründung'),
    ('Beurteilung bestätigt', 'Koordination / Azubi', 'Praxisbeurteilung abgeschlossen'),
    ('Ausbildungsnachweis genehmigt', 'Azubi', 'Wochenbericht freigegeben'),
    ('Ausbildungsnachweis abgelehnt', 'Azubi', 'Mit Korrekturhinweis'),
    ('Nachweis-Erinnerung', 'Azubi', 'Fehlende/überfällige Nachweise'),
    ('Portal-Willkommen', 'Azubi', 'Zugangsdaten für das Portal'),
    ('Lerntag genehmigt', 'Azubi', 'Antrag freigegeben'),
    ('Lerntag abgelehnt', 'Azubi', 'Antrag abgelehnt'),
    ('Urlaub genehmigt', 'Azubi', 'Urlaubsantrag freigegeben'),
    ('Urlaub abgelehnt', 'Azubi', 'Urlaubsantrag abgelehnt'),
    ('Urlaub bearbeitet', 'Azubi', 'Durch Urlaubsstelle abgeschlossen'),
    ('Inventar ausgegeben', 'Azubi', 'Gegenstand erhalten'),
    ('Inventar zurückgegeben', 'Azubi', 'Rückgabe bestätigt'),
    ('Beurteilungstoken versendet', 'Praxistutor', 'Link zur Online-Beurteilung'),
    ('Beurteilung eingereicht', 'Koordination', 'Praxistutor hat ausgefüllt'),
]
for row in notifications:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('15.2  Vorlagen anpassen')

para(
    'E-Mail-Vorlagen werden unter Systemkonfiguration → Benachrichtigungsvorlagen '
    'verwaltet. Jede Vorlage hat:'
)
bullet('Betreffzeile mit Platzhalterunterstützung')
bullet('HTML-Textkörper mit variablen Feldern ({{ student_vorname }}, {{ von }}, etc.)')
bullet('Aktiviert / Deaktiviert-Schalter')

h2('15.3  Interne Benachrichtigungen')

para(
    'Zusätzlich zur E-Mail erhalten Mitarbeiter In-App-Benachrichtigungen '
    '(Glockensymbol in der Navigationsleiste). '
    'Diese können als gelesen markiert werden.'
)

h2('15.4  Benachrichtigungseinstellungen pro Nutzer')

para(
    'Jeder Nutzer kann im Profil bestimmte E-Mail-Benachrichtigungen deaktivieren '
    '(z. B. Nachweis-Erinnerungen, Einsatz-Erinnerungen).'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 16: DOKUMENTE & VORLAGEN
# ═══════════════════════════════════════════════════════════════════════════════

h1('16  Dokumente & Vorlagen')

h2('16.1  Dokumentenvorlagen (.docx)')

para(
    'Das System generiert Dokumente auf Basis von Word-Vorlagen (.docx). '
    'Vorlagen verwenden Jinja2-Syntax für Platzhalter: {{ vorname }}, {{ datum }}, etc. '
    'Vorlagen werden im jeweiligen Modul unter Einstellungen hochgeladen.'
)

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Dokumenttyp', 'Verfügbare Platzhalter (Auswahl)'], COLOR_DARK)
doc_types = [
    ('Zuweisungsschreiben', 'student_vorname, student_nachname, einheit_name, block_beginn, block_ende'),
    ('Praktikumsplan', 'student_*, kurs_*, liste aller Einsätze als Schleife'),
    ('Stationsbrief', 'student_*, station_*, von, bis, praxistutor_*'),
    ('Nachwuchskraft-Dokument', 'Alle Stammdaten, kurs_*, status_*'),
    ('Reservierungsbestätigung', 'student_*, zimmer_*, wohnheim_*, von, bis'),
    ('Urlaubsbestätigung', 'Liste aller Anträge: vorname, nachname, von, bis, arbeitstage'),
    ('Inventarquittung', 'student_*, artikel_*, kategorie_*, ausgegeben_am'),
    ('Ausbildungsnachweis-Export', 'Alle Wochennachweise mit Tageseinträgen'),
]
for row in doc_types:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('16.2  Signierte PDFs (Elektronische Signatur)')

para(
    'Für Urlaubsgenehmigungen und Beurteilungsbestätigungen können signierte PDFs '
    'heruntergeladen werden. Das Dokument enthält:'
)
bullet('Alle relevanten Dokumentdaten')
bullet('Elektronischen Signaturblock: Name, Funktion, Datum/Uhrzeit')
bullet('Hinweis: Einfache elektronische Signatur gem. eIDAS Art. 3 Nr. 10')

info_box(
    'Signierte PDFs werden bei der Genehmigung automatisch vorbereitet. '
    'Ein Download-Button erscheint nach der Genehmigung/Bestätigung.'
)

h2('16.3  Paperless-ngx Integration')

para(
    'Das System kann optional mit einer Paperless-ngx-Instanz verbunden werden. '
    'Generierte Dokumente werden dann automatisch dort archiviert und '
    'mit dem jeweiligen Azubi oder Kurs als Korrespondent verknüpft. '
    'Im System können Dokumente direkt aus Paperless vorschau angezeigt werden.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 17: ORGANISATION
# ═══════════════════════════════════════════════════════════════════════════════

h1('17  Organisation')

h2('17.1  Organisationseinheiten')

para(
    'Die Organisationsstruktur bildet die Aufbauorganisation des Unternehmens/Behörde ab. '
    'Sie dient als Grundlage für Praktikumseinsätze und Kapazitätsplanung.'
)

para('Hierarchieebenen (von oben nach unten):')
numbered('Behörde / Konzern')
numbered('Abteilung')
numbered('Referatsgruppe')
numbered('Referat')
numbered('Sachgebiet')

para('Jede Einheit hat:')
bullet('Name und Kürzel')
bullet('Übergeordnete Einheit (Parent)')
bullet('Zugehöriger Standort')
bullet('Kapazitätsgrenze (max. gleichzeitige Azubis)')

h2('17.2  Standorte')

para(
    'Standorte (Gebäude, Liegenschaften) werden mit vollständiger Adresse gepflegt '
    'und können Organisationseinheiten zugeordnet werden.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 18: PRAXISTUTOREN & KOORDINATIONSGRUPPEN
# ═══════════════════════════════════════════════════════════════════════════════

h1('18  Praxistutoren & Ausbildungskoordination')

h2('18.1  Praxistutoren anlegen')

para('Praxistutoren sind externe Betreuer in den Praktikumseinheiten:')
bullet('Vorname, Nachname, E-Mail-Adresse')
bullet('Arbeitsort / Organisationseinheit')
bullet('Zugeordnete Berufsbilder')
bullet('Status: Ausstehend / Bestätigt')

para(
    'Nach Bestätigung erhält der Praxistutor eine Willkommens-E-Mail. '
    'Für Beurteilungen wird kein eigenes Benutzerkonto benötigt – '
    'es wird ein Token-Link per E-Mail zugesandt.'
)

h2('18.2  Koordinationsgruppen')

para(
    'Koordinationsgruppen bündeln mehrere Mitarbeiter, '
    'die gemeinsam für bestimmte Organisationseinheiten und Berufsbilder '
    'zuständig sind:'
)
bullet('Mehrere Mitglieder pro Gruppe')
bullet('Funktionspostfach (gemeinsame E-Mail-Adresse für Gruppen-Mails)')
bullet('Zugeordnete Organisationseinheiten')
bullet('Portal-Zugänge für Koordinationsmitglieder')

h2('18.3  Chief Instructors')

para(
    'Der Chief Instructor leitet eine Koordinationsgruppe und ist berechtigt, '
    'Praktikumseinsätze zu genehmigen oder abzulehnen.'
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 19: SYSTEMKONFIGURATION & ADMINISTRATION
# ═══════════════════════════════════════════════════════════════════════════════

h1('19  Systemkonfiguration & Administration')

h2('19.1  Globale Einstellungen')

para('Unter Einstellungen → Systemkonfiguration (nur Ausbildungsleitung):')

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Einstellung', 'Beschreibung'], COLOR_DARK)
config_items = [
    ('Markenname', 'Wird im Portal-Header und E-Mails angezeigt'),
    ('Primärfarbe / Sekundärfarbe', 'Brandingfarben des Portals'),
    ('Impressum / Datenschutztexte', 'Rechtliche Pflichtseiten'),
    ('Erinnerung Einsatzbeginn/-ende', 'Vorlauf in Tagen (Standard: 7)'),
    ('Erinnerungszeit', 'Uhrzeit der täglichen Erinnerungen (Standard: 07:00)'),
    ('Anonymisierungsfrist', 'Monate Inaktivität bis zur Anonymisierung (Standard: 12)'),
    ('Anonymisierungsuhrzeit', 'Täglicher Zeitpunkt der Anonymisierung'),
]
for row in config_items:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

h2('19.2  Django-Adminpanel')

para(
    'Das Django-Adminpanel ist unter /admin/ erreichbar '
    '(nur für Superuser und Staff-Benutzer). Hier werden verwaltet:'
)
bullet('Benutzerkonten und Berechtigungen')
bullet('Rollenprofile (Ausbildungsreferat-Profile)')
bullet('Statussystem, Berufsbilder, Beschäftigungsverhältnisse')
bullet('Benachrichtigungsvorlagen')
bullet('Audit-Log (vollständige Änderungshistorie)')
bullet('Alle Stammdaten-Tabellen')

h2('19.3  Hintergrundjobs (Celery)')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Task', 'Uhrzeit', 'Funktion'], COLOR_DARK)
tasks = [
    ('Einsatz-Erinnerungen', '07:00 Uhr', 'E-Mail 7 Tage vor Einsatzbeginn/-ende'),
    ('Nachweis-Erinnerungen', '07:00 Uhr', 'Erinnerung an fehlende Ausbildungsnachweise'),
    ('Urlaubspaket-Versand', '08:00 Uhr', 'Genehmigte Anträge an Urlaubsstelle'),
    ('Krankmeldungs-Report', '08:05 Uhr', 'Tagesbericht an Urlaubsstelle'),
    ('Anonymisierung', '03:00 Uhr', 'Automatische Anonymisierung inaktiver Datensätze'),
]
for row in tasks:
    add_row(t, *row)
style_table(t)
doc.add_paragraph()

para('Status der Hintergrundjobs überwachen:')
code('docker compose logs -f celery_worker')
code('docker compose logs -f celery_beat')

h2('19.4  Audit-Log')

para(
    'Das System protokolliert alle Datenänderungen automatisch im Audit-Log. '
    'Für jede Änderung werden erfasst:'
)
bullet('Zeitstempel (Datum und Uhrzeit)')
bullet('Benutzer (wer hat geändert)')
bullet('Aktion: Erstellt / Geändert / Gelöscht')
bullet('Betroffenes Objekt und veränderte Felder')
bullet('Optionale Verknüpfung mit einer Nachwuchskraft')

para('Das Audit-Log ist im Adminpanel einsehbar und nicht löschbar.')

h2('19.5  Sicherheitshinweise')

para('Für den sicheren Betrieb sind folgende Punkte zu beachten:')
bullet('DEBUG=False im Produktivbetrieb (niemals True auf einem öffentlichen Server)')
bullet('SECRET_KEY mindestens 50 zufällige Zeichen, regelmäßig rotieren')
bullet('ALLOWED_HOSTS auf tatsächlich verwendete Hostnamen beschränken')
bullet('Datenbankpasswort und PAPERLESS_API_KEY sicher verwahren (.env nie ins Git!)')
bullet('HTTPS/TLS vor dem System (z. B. nginx als Reverse Proxy mit Let\'s Encrypt)')
bullet('Regelmäßige Backups der PostgreSQL-Datenbank und des Medien-Volumes')
bullet('Container-Images regelmäßig aktualisieren (docker compose pull)')


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# KAPITEL 20: HÄUFIGE FRAGEN & TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════

h1('20  Häufige Fragen & Troubleshooting')

h2('20.1  Häufige Fragen')

h3('Eine Nachwuchskraft kann sich nicht ins Portal einloggen.')
para(
    'Prüfen, ob dem Azubi-Datensatz ein Portalbenutzer zugeordnet ist. '
    'Im Admin unter „Student → Nachwuchskräfte" den Datensatz öffnen und '
    'das Feld „Portal-Benutzer" prüfen. Ggf. Benutzerkonto anlegen und '
    'Willkommens-E-Mail versenden.'
)

h3('E-Mails werden nicht versendet.')
para(
    'SMTP-Einstellungen in der .env prüfen (EMAIL_HOST, EMAIL_PORT, Passwort, TLS). '
    'Test über das Django-Adminpanel: „E-Mail-Test" (falls konfiguriert). '
    'Celery-Worker-Logs prüfen: docker compose logs celery_worker.'
)

h3('Celery-Tasks laufen nicht.')
para(
    'Prüfen, ob Redis erreichbar ist (docker compose ps). '
    'Celery-Worker neu starten: docker compose restart celery_worker celery_beat.'
)

h3('Die Paperless-Integration liefert keine Dokumente.')
para(
    'PAPERLESS_URL und PAPERLESS_API_KEY in der .env prüfen. '
    'Netzwerkverbindung zwischen den Containern testen. '
    'Paperless-ngx-Logs prüfen.'
)

h3('Ein Benutzer sieht die falsche Seite oder bekommt „403 Forbidden".')
para(
    'Rollenprofile im Admin prüfen. Ggf. Cache leeren: '
    'docker compose exec app python manage.py shell -c "from django.core.cache import cache; cache.clear()"'
)

h2('20.2  Logs einsehen')

code('# Alle Container-Logs')
code('docker compose logs -f')
code('')
code('# Nur App-Logs')
code('docker compose logs -f app')
code('')
code('# Nur Celery-Logs')
code('docker compose logs -f celery_worker celery_beat')

h2('20.3  Neustart einzelner Dienste')

code('docker compose restart app')
code('docker compose restart celery_worker')
code('docker compose restart celery_beat')


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ANHANG
# ═══════════════════════════════════════════════════════════════════════════════

h1('Anhang: Glossar')

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
add_table_header_row(t, ['Begriff', 'Erläuterung'], COLOR_DARK)
glossary = [
    ('Nachwuchskraft', 'Auszubildende/r im System (Azubi)'),
    ('Berufsbild', 'Ausbildungsberuf mit allen zugehörigen Konfigurationen'),
    ('Kurs', 'Gruppe von Nachwuchskräften desselben Berufsbilds und Jahrgangs'),
    ('Ablaufblock', 'Phase innerhalb eines Kurses (Theorie, Praktikum, etc.)'),
    ('Praktikumseinsatz', 'Konkrete Zuweisung einer Nachwuchskraft an eine OE für einen Block'),
    ('OE', 'Organisationseinheit (Abteilung, Referat, etc.)'),
    ('Praxistutor', 'Betreuer in der aufnehmenden Organisationseinheit'),
    ('Chief Instructor', 'Leiter einer Ausbildungskoordinations-Gruppe'),
    ('Koordination', 'Ausbildungskoordinationsgruppe (Team aus Chief + Mitgliedern)'),
    ('Ausbildungsnachweis', 'Wöchentlicher Bericht der Nachwuchskraft (Berichtsheft)'),
    ('Abwesenheitsampel', 'Automatische Fehlzeitenquote (grün/gelb/rot)'),
    ('Lerntag', 'Selbststudientag außerhalb des regulären Dienstplans'),
    ('Token', 'Einmaliger, sicherer Zugriffslink (z. B. für Praxistutoren oder Urlaubsstelle)'),
    ('Paperless-ngx', 'Externes Open-Source-Dokumenten-Management-System'),
    ('Celery', 'Python-Framework für Hintergrundaufgaben und Zeitplanung'),
    ('eIDAS', 'EU-Verordnung über elektronische Identifizierung und Vertrauensdienste'),
    ('einfache e. Signatur', 'Identifikation einer Person durch Zeitstempel und Name (eIDAS Art. 3 Nr. 10)'),
]
for row in glossary:
    add_row(t, *row)
style_table(t)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f'Azubi-Portal Handbuch  ·  Version 1.0  ·  '
    f'Stand {datetime.date.today().strftime("%d.%m.%Y")}'
)
run.font.color.rgb = COLOR_GRAY
run.font.size = Pt(9)
run.font.italic = True

# ── Speichern ─────────────────────────────────────────────────────────────────
import os
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Azubi_Portal_Handbuch.docx')
doc.save(output_path)
print(f'Handbuch gespeichert: {output_path}')
